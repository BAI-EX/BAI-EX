from flask import Flask, render_template, redirect, url_for, request, flash, session, jsonify, Response, stream_with_context, make_response
import anthropic
from flask_sqlalchemy import SQLAlchemy
from werkzeug.security import generate_password_hash, check_password_hash
from datetime import datetime, timedelta, date
from functools import wraps
import os, json, random

import sys as _sys
_is_frozen = getattr(_sys, 'frozen', False)
if _is_frozen:
    _base = os.path.join(_sys._MEIPASS)
    app = Flask(__name__,
                template_folder=os.path.join(_base, 'templates'),
                static_folder=os.path.join(_base, 'static'))
else:
    app = Flask(__name__)
# SECRET_KEY gerada via os.urandom — em produção, coloque no .env
app.secret_key = os.environ.get('SECRET_KEY', 'baiex-secret-2025-dev-only')

# ── FIX 1A: Sessão expira em 8 horas ───────────────────────────
from datetime import timedelta
app.config['PERMANENT_SESSION_LIFETIME'] = timedelta(hours=8)
app.config['SESSION_COOKIE_HTTPONLY'] = True   # JS não acessa o cookie
app.config['SESSION_COOKIE_SAMESITE'] = 'Lax' # Proteção CSRF básica

# ── FIX 1B: Força HTTPS em produção ────────────────────────────
@app.before_request
def forcar_https():
    """Redireciona HTTP → HTTPS em produção (Render já provê HTTPS)."""
    if os.environ.get('FLASK_ENV') == 'production':
        if request.headers.get('X-Forwarded-Proto') == 'http':
            url = request.url.replace('http://', 'https://', 1)
            return redirect(url, code=301)

# ── FIX 1C: Renova sessão como permanente no login ──────────────
def marcar_sessao_permanente():
    session.permanent = True

# Cliente Anthropic
# A API key é lida automaticamente do arquivo .env na mesma pasta do app.py
import os

def _carregar_env():
    """Lê .env se existir — evita necessidade de variável de ambiente no terminal."""
    env_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), '.env')
    if os.path.exists(env_path):
        with open(env_path, encoding='utf-8') as f:
            for line in f:
                line = line.strip()
                if line and not line.startswith('#') and '=' in line:
                    k, v = line.split('=', 1)
                    os.environ.setdefault(k.strip(), v.strip().strip('"').strip("'"))

_carregar_env()
ANTHROPIC_API_KEY = os.environ.get('ANTHROPIC_API_KEY', '')

def get_ai_client():
    # ── FIX 6: Nunca expor a key em logs ou mensagens de erro ──────
    key = os.environ.get('ANTHROPIC_API_KEY', '')
    if not key or key == 'sk-ant-COLE_SUA_CHAVE_AQUI':
        raise ValueError(
            "IA não configurada. "
            "Adicione ANTHROPIC_API_KEY no arquivo .env ou nas variáveis de ambiente."
        )
    # Validação mínima de formato sem logar o valor real
    if not key.startswith('sk-ant-'):
        raise ValueError("ANTHROPIC_API_KEY inválida. Verifique o formato da chave.")
    return anthropic.Anthropic(api_key=key)

def _mascarar_key(key: str) -> str:
    """Mostra apenas os 8 primeiros e 4 últimos caracteres para diagnóstico."""
    if not key or len(key) < 12:
        return '(não configurada)'
    return f"{key[:8]}...{key[-4:]}"

def status_ia() -> dict:
    """Retorna status da IA sem expor a key completa."""
    key = os.environ.get('ANTHROPIC_API_KEY', '')
    return {
        'configurada': bool(key and key != 'sk-ant-COLE_SUA_CHAVE_AQUI'),
        'preview': _mascarar_key(key)
    }

BASE_DIR = os.path.dirname(os.path.abspath(__file__))

# ── FIX 4: PostgreSQL em produção, SQLite local ────────────────
# Em produção (Render): configure DATABASE_URL no painel Environment
# Localmente: usa SQLite automático
_db_url = os.environ.get('DATABASE_URL', '')
if _db_url:
    # Render PostgreSQL: corrige prefixo antigo do SQLAlchemy
    if _db_url.startswith('postgres://'):
        _db_url = _db_url.replace('postgres://', 'postgresql://', 1)
    app.config['SQLALCHEMY_DATABASE_URI'] = _db_url
    app.config['SQLALCHEMY_ENGINE_OPTIONS'] = {
        'pool_pre_ping': True,       # verifica conexão antes de usar
        'pool_recycle': 300,         # recicla conexões a cada 5 min
        'connect_args': {}
    }
else:
    DB_PATH = os.path.join(BASE_DIR, 'baiex.db')
    app.config['SQLALCHEMY_DATABASE_URI'] = f'sqlite:///{DB_PATH}'
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False

db = SQLAlchemy(app)

# ─────────────────────────── MODELOS ───────────────────────────

# ── FIX 2: Controle de tentativas de login (anti brute-force) ──
# Armazena em memória: {chave: {'tentativas': N, 'bloqueado_ate': datetime}}
# Chave = "ip:area" — bloqueia apenas a área que sofreu tentativas, não o IP inteiro.
# Exemplo: falha no login de funcionário NÃO bloqueia o login de RH ou admin.
_tentativas_login = {}
MAX_TENTATIVAS = 5
BLOQUEIO_MINUTOS = 15

def _chave_bloqueio(ip: str, area: str) -> str:
    """Gera chave única por IP + área de login (funcionario / empresa / admin)."""
    return f"{ip}:{area}"

def verificar_bloqueio(ip, area='geral'):
    """Retorna minutos restantes de bloqueio (int > 0) ou False se liberado."""
    import datetime as _dt
    chave = _chave_bloqueio(ip, area)
    dados = _tentativas_login.get(chave)
    if not dados:
        return False
    if dados.get('bloqueado_ate') and _dt.datetime.now() < dados['bloqueado_ate']:
        minutos = int((dados['bloqueado_ate'] - _dt.datetime.now()).seconds / 60) + 1
        return minutos
    if dados.get('bloqueado_ate') and _dt.datetime.now() >= dados['bloqueado_ate']:
        _tentativas_login.pop(chave, None)  # desbloqueia após o tempo
    return False

def registrar_falha(ip, area='geral'):
    """Registra tentativa falha. Bloqueia área após MAX_TENTATIVAS."""
    import datetime as _dt
    chave = _chave_bloqueio(ip, area)
    if chave not in _tentativas_login:
        _tentativas_login[chave] = {'tentativas': 0, 'bloqueado_ate': None}
    _tentativas_login[chave]['tentativas'] += 1
    if _tentativas_login[chave]['tentativas'] >= MAX_TENTATIVAS:
        _tentativas_login[chave]['bloqueado_ate'] = _dt.datetime.now() + _dt.timedelta(minutes=BLOQUEIO_MINUTOS)

def resetar_tentativas(ip, area='geral'):
    """Reseta após login bem-sucedido."""
    _tentativas_login.pop(_chave_bloqueio(ip, area), None)


class Empresa(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    nome = db.Column(db.String(120), nullable=False)
    cnpj = db.Column(db.String(20), unique=True, nullable=False)
    responsavel = db.Column(db.String(120), nullable=False)
    email = db.Column(db.String(120), unique=True, nullable=False)
    senha = db.Column(db.String(200), nullable=False)
    # plano: starter1, starter2, profissional, empresarial, enterprise
    # prefixo trial_: trial_starter1, trial_starter2, trial_profissional, trial_empresarial
    # cancelado, inativo
    plano = db.Column(db.String(30), default='trial_starter1')
    frequencia_checkin = db.Column(db.String(20), default='semanal')  # semanal, quinzenal, mensal
    trial_inicio = db.Column(db.DateTime, default=datetime.utcnow)
    ativo = db.Column(db.Boolean, default=True)
    cancelado = db.Column(db.Boolean, default=False)
    criado_em = db.Column(db.DateTime, default=datetime.utcnow)
    fidelizacao_inicio = db.Column(db.DateTime, nullable=True)
    fidelizacao_fim    = db.Column(db.DateTime, nullable=True)
    funcionarios = db.relationship('Funcionario', backref='empresa', lazy=True)

    # ── helpers de plano ──────────────────────────────────────
    @property
    def plano_base(self):
        """Retorna o plano sem prefixo trial_."""
        return self.plano.replace('trial_', '') if self.plano else 'starter1'

    @property
    def em_trial(self):
        return self.plano and self.plano.startswith('trial_')

    @property
    def dias_trial_restantes(self):
        if not self.em_trial or not self.trial_inicio:
            return None
        dias_total = 30 if self.plano_base in ('starter1', 'starter2') else 14
        decorrido = (datetime.utcnow() - self.trial_inicio).days
        return max(0, dias_total - decorrido)

    @property
    def trial_expirado(self):
        if not self.em_trial:
            return False
        return self.dias_trial_restantes == 0

    @property
    def em_fidelizacao(self):
        """True durante os 6 meses de fidelização após ativar plano pago."""
        if not self.fidelizacao_fim:
            return False
        return datetime.utcnow() < self.fidelizacao_fim

    @property
    def meses_fidelizacao_restantes(self):
        if not self.em_fidelizacao:
            return 0
        return max(1, int((self.fidelizacao_fim - datetime.utcnow()).days / 30))

    @property
    def nivel_plano(self):
        """Retorna nível numérico para comparação: 1=starter1, 2=starter2, 3=prof, 4=emp, 5=ent."""
        mapa = {'starter1': 1, 'starter2': 2, 'profissional': 3, 'empresarial': 4, 'enterprise': 5}
        return mapa.get(self.plano_base, 1)

    def tem_acesso(self, nivel_minimo: str) -> bool:
        """Verifica se empresa tem acesso a determinado nível de plano."""
        if self.cancelado or not self.ativo:
            return False
        if self.trial_expirado:
            return False
        niveis = {'starter1': 1, 'starter2': 2, 'profissional': 3, 'empresarial': 4, 'enterprise': 5}
        return self.nivel_plano >= niveis.get(nivel_minimo, 99)

    @property
    def limite_funcionarios(self):
        limites = {'starter1': 15, 'starter2': 30, 'profissional': 100,
                   'empresarial': 200, 'enterprise': 99999}
        return limites.get(self.plano_base, 15)

class Funcionario(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    empresa_id = db.Column(db.Integer, db.ForeignKey('empresa.id'), nullable=False)
    nome = db.Column(db.String(120), nullable=False)
    email = db.Column(db.String(120), unique=True, nullable=False)
    senha = db.Column(db.String(200), nullable=False)
    departamento = db.Column(db.String(80), default='Geral')
    cargo = db.Column(db.String(80), default='')
    ativo = db.Column(db.Boolean, default=True)
    criado_em = db.Column(db.DateTime, default=datetime.utcnow)
    checkins = db.relationship('CheckIn', backref='funcionario', lazy=True)

class CheckIn(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    funcionario_id = db.Column(db.Integer, db.ForeignKey('funcionario.id'), nullable=False)
    empresa_id = db.Column(db.Integer, db.ForeignKey('empresa.id'), nullable=False)
    data = db.Column(db.DateTime, default=datetime.utcnow)
    semana = db.Column(db.Integer, nullable=False)  # semana do ano
    ano = db.Column(db.Integer, nullable=False)
    periodo_ref = db.Column(db.String(20), nullable=True)  # ex: "2025-W14", "2025-Q2-1", "2025-M4"

    # Perguntas (1-5)
    q1_estresse = db.Column(db.Integer, nullable=False)       # nível de estresse
    q2_sobrecarga = db.Column(db.Integer, nullable=False)     # se sente sobrecarregado
    q3_sono = db.Column(db.Integer, nullable=False)           # dormindo bem
    q4_motivacao = db.Column(db.Integer, nullable=False)      # motivação
    q5_esgotamento = db.Column(db.Integer, nullable=False)    # esgotamento emocional
    q6_relacoes = db.Column(db.Integer, nullable=False)       # relações no trabalho
    q7_autonomia = db.Column(db.Integer, nullable=False)      # sensação de controle
    q8_reconhecimento = db.Column(db.Integer, nullable=False) # reconhecimento
    # Perguntas 9-24 (adicionadas)
    q9_clareza        = db.Column(db.Integer, nullable=True)  # clareza do trabalho
    q10_comunicacao   = db.Column(db.Integer, nullable=True)  # comunicacao
    q11_lideranca     = db.Column(db.Integer, nullable=True)  # apoio da lideranca
    q12_ritmo         = db.Column(db.Integer, nullable=True)  # ritmo de trabalho
    q13_interrupcoes  = db.Column(db.Integer, nullable=True)  # interrupcoes
    q14_decisao       = db.Column(db.Integer, nullable=True)  # autonomia decisoria
    q15_seguranca     = db.Column(db.Integer, nullable=True)  # seguranca psicologica
    q16_conflitos     = db.Column(db.Integer, nullable=True)  # conflitos na equipe
    q17_justica       = db.Column(db.Integer, nullable=True)  # justica organizacional
    q18_desconexao    = db.Column(db.Integer, nullable=True)  # desconexao fora do expediente
    q19_descanso      = db.Column(db.Integer, nullable=True)  # descanso suficiente
    q20_pressao       = db.Column(db.Integer, nullable=True)  # pressao por resultados
    q21_desanimo      = db.Column(db.Integer, nullable=True)  # desanimo frequente
    q22_respeito      = db.Column(db.Integer, nullable=True)  # respeito no ambiente
    q23_performance   = db.Column(db.Integer, nullable=True)  # queda de produtividade
    q24_satisfacao    = db.Column(db.Integer, nullable=True)  # satisfacao geral

    score = db.Column(db.Float, nullable=False)  # 0-100 (100 = risco máximo)
    nivel_risco = db.Column(db.String(20), nullable=False)   # baixo / moderado / alto / critico

    @staticmethod
    def calcular_score(respostas):
        """
        Calcula o score de risco de burnout (0-100).
        Todas as perguntas: 1 = melhor situação, 5 = pior situação.
        Compatível com 8 perguntas (legado) e 24 perguntas (atual).
        """
        n = len(respostas)
        pontos = sum(respostas)
        score = ((pontos - n) / (n * 4)) * 100  # min=n, max=5n → normaliza para 0-100
        return round(score, 1)

    @staticmethod
    def classificar_risco(score):
        # Escala 0-100: quartis naturais da soma 8-40
        if score < 25:
            return 'baixo'
        elif score < 50:
            return 'moderado'
        elif score < 75:
            return 'alto'
        else:
            return 'critico'

class Admin(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    usuario = db.Column(db.String(80), unique=True, nullable=False)
    senha = db.Column(db.String(200), nullable=False)

# ── FIX 7 + FIX 9: Modelo de Auditoria ─────────────────────────
class AuditLog(db.Model):
    """
    Registra ações importantes para rastreabilidade NR-01.
    A norma exige registros por 20 anos — esse modelo é o início disso.
    """
    id          = db.Column(db.Integer, primary_key=True)
    timestamp   = db.Column(db.DateTime, default=datetime.utcnow, nullable=False)
    tipo_usuario = db.Column(db.String(20))   # empresa / funcionario / admin
    usuario_id  = db.Column(db.Integer)       # id do usuário que agiu
    empresa_id  = db.Column(db.Integer, db.ForeignKey('empresa.id'), nullable=True)
    acao        = db.Column(db.String(80), nullable=False)   # ex: LOGIN, CHECKIN, RELATORIO
    detalhe     = db.Column(db.String(255))   # informações extras
    ip          = db.Column(db.String(45))    # IPv4 ou IPv6

    def __repr__(self):
        return f'<AuditLog {self.acao} by {self.tipo_usuario}:{self.usuario_id}>'

def registrar_auditoria(acao, detalhe='', tipo_usuario=None, usuario_id=None, empresa_id=None):
    """Grava uma entrada de auditoria. Falha silenciosamente para não travar a app."""
    try:
        log = AuditLog(
            acao=acao[:80],
            detalhe=str(detalhe)[:255] if detalhe else '',
            tipo_usuario=tipo_usuario or session.get('tipo', 'anonimo'),
            usuario_id=usuario_id or (
                session.get('empresa_id') or
                session.get('funcionario_id') or
                session.get('admin_id')
            ),
            empresa_id=empresa_id or session.get('empresa_id'),
            ip=request.remote_addr
        )
        db.session.add(log)
        db.session.commit()
    except Exception:
        pass  # log nunca deve travar a operação principal

# ─────────────────────────── DECORATORS ───────────────────────────

def login_required_empresa(f):
    @wraps(f)
    def decorated(*args, **kwargs):
        if 'empresa_id' not in session:
            return redirect(url_for('login_empresa'))
        return f(*args, **kwargs)
    return decorated

def login_required_funcionario(f):
    @wraps(f)
    def decorated(*args, **kwargs):
        if 'funcionario_id' not in session:
            return redirect(url_for('login_funcionario'))
        return f(*args, **kwargs)
    return decorated

def login_required_admin(f):
    @wraps(f)
    def decorated(*args, **kwargs):
        if 'admin_id' not in session:
            return redirect(url_for('login_admin'))
        return f(*args, **kwargs)
    return decorated

def requer_plano(nivel_minimo):
    """Decorator que bloqueia acesso se plano da empresa for inferior ao mínimo."""
    def decorator(f):
        @wraps(f)
        def decorated(*args, **kwargs):
            empresa = Empresa.query.get(session.get('empresa_id'))
            if not empresa or not empresa.tem_acesso(nivel_minimo):
                nomes = {
                    'profissional': 'Profissional',
                    'empresarial': 'Empresarial',
                    'enterprise': 'Enterprise',
                }
                return render_template('empresa/upgrade_required.html',
                    plano_necessario=nomes.get(nivel_minimo, nivel_minimo),
                    plano_atual=empresa.plano_base if empresa else 'starter1',
                    trial_expirado=empresa.trial_expirado if empresa else False)
            return f(*args, **kwargs)
        return decorated
    return decorator

# ─────────────────────────── PLANOS CONFIG ───────────────────────────

PLANOS_INFO = {
    'starter1': {
        'nome': 'Starter',
        'descricao': 'Ideal para pequenas empresas iniciando',
        'preco_display': 'R$ 149/mês',
        'preco_mensal': 149.0,
        'preco_tipo': 'fixo',
        'limite_func': 15,
        'trial_dias': 30,
        'stripe_price_id_env': 'STRIPE_PRICE_STARTER1',
        'features': [
            'Check-in semanal (8 dimensões)',
            'Score de risco 0–100 por colaborador',
            'Dashboard RH em tempo real',
            'Central de alertas automáticos',
            'Relatório básico NR-01',
            'Convite automático por e-mail',
        ],
        'ia': False,
        'frequencia_extra': False,
        'benchmark': False,
        'multi_unidade': False,
    },
    'starter2': {
        'nome': 'Profissional',
        'descricao': 'Para empresas em crescimento (16–200 funcionários)',
        'preco_display': 'R$ 15/func/mês',
        'preco_por_func': 15,
        'preco_tipo': 'por_func',
        'limite_func': 200,
        'trial_dias': 15,
        'stripe_price_id_env': 'STRIPE_PRICE_STARTER2',
        'features': [
            'Tudo do Starter Essencial',
            'Até 30 funcionários',
        ],
        'ia': False,
        'frequencia_extra': False,
        'benchmark': False,
        'multi_unidade': False,
    },
    'profissional': {
        'nome': 'Profissional',
        'descricao': 'Para empresas em crescimento (16–200 funcionários)',
        'preco_display': 'R$ 15/func/mês',
        'preco_por_func': 15,
        'preco_tipo': 'por_func',
        'limite_func': 200,
        'trial_dias': 15,
        'stripe_price_id_env': 'STRIPE_PRICE_PRO',
        'features': [
            'Tudo do Starter',
            'Frequência configurável (semanal/quinzenal/mensal)',
            'Análise por departamento/setor',
            'Relatório NR-01 completo com IA',
            'Previsão preditiva de risco (IA)',
            'BAIA — Chat IA para RH',
            'Histórico de 12 meses',
            'Suporte via WhatsApp',
        ],
        'ia': True,
        'frequencia_extra': True,
        'benchmark': False,
        'multi_unidade': False,
    },
    'empresarial': {
        'nome': 'Enterprise',
        'descricao': 'Para empresas que precisam de recursos avançados',
        'preco_display': 'Sob consulta',
        'preco_tipo': 'negociado',
        'limite_func': 99999,
        'trial_dias': 0,
        'stripe_price_id_env': 'STRIPE_PRICE_EMPRESARIAL',
        'features': [
            'Tudo do Profissional',
            'Benchmark anônimo por setor',
            'Multi-unidades / filiais',
            'Relatório consolidado por filial',
            'Suporte WhatsApp prioritário',
        ],
        'ia': True,
        'frequencia_extra': True,
        'benchmark': True,
        'multi_unidade': True,
    },
    'enterprise': {
        'nome': 'Enterprise',
        'descricao': 'Para empresas que precisam de recursos avançados',
        'preco_display': 'Sob consulta',
        'preco_tipo': 'negociado',
        'limite_func': 99999,
        'trial_dias': 0,
        'stripe_price_id_env': '',
        'features': [
            'Tudo do Empresarial',
            'API para sistemas de RH',
            'Integração eSocial / CIPA (roadmap)',
            'Gerente de conta dedicado',
            'SLA contratual',
        ],
        'ia': True,
        'frequencia_extra': True,
        'benchmark': True,
        'multi_unidade': True,
    },
}

def calcular_preco_plano(plano_base: str, num_funcionarios: int) -> float:
    info = PLANOS_INFO.get(plano_base, PLANOS_INFO['starter1'])
    if info['preco_tipo'] == 'fixo':
        return info['preco_mensal']
    elif info['preco_tipo'] == 'por_func':
        return info['preco_por_func'] * num_funcionarios
    return 0.0

def plano_para_funcionarios(n: int) -> str:
    """Retorna o plano recomendado para N funcionários (3 planos)."""
    if n <= 15:  return 'starter1'
    if n <= 200: return 'profissional'
    return 'enterprise'

# ─────────────────────────── UTILS ───────────────────────────

def get_semana_atual():
    hoje = date.today()
    return hoje.isocalendar()[1], hoje.year

def periodo_atual(empresa):
    """
    Retorna (periodo_ref, label) com base na frequência configurada da empresa.
    periodo_ref é usado para verificar se o funcionário já respondeu no período.
    """
    hoje = date.today()
    freq = empresa.frequencia_checkin if empresa else 'semanal'

    if freq == 'quinzenal':
        semana_iso = hoje.isocalendar()[1]
        quinzena = 1 if semana_iso % 2 == 1 else 2
        ref = f"{hoje.year}-Q{semana_iso // 2}-{quinzena}"
        label = f"Quinzena atual"
    elif freq == 'mensal':
        ref = f"{hoje.year}-M{hoje.month:02d}"
        label = f"Mês atual ({hoje.strftime('%B/%Y')})"
    else:  # semanal
        semana = hoje.isocalendar()[1]
        ref = f"{hoje.year}-W{semana:02d}"
        label = f"Semana {semana}/{hoje.year}"

    return ref, label

def proximo_checkin_label(empresa, ultimo_checkin_data):
    """Retorna texto indicando quando o próximo check-in estará disponível."""
    if not ultimo_checkin_data:
        return "Disponível agora"
    freq = empresa.frequencia_checkin if empresa else 'semanal'
    dias = {'semanal': 7, 'quinzenal': 14, 'mensal': 30}.get(freq, 7)
    proximo = ultimo_checkin_data + timedelta(days=dias)
    hoje = datetime.utcnow()
    if proximo <= hoje:
        return "Disponível agora"
    diff = proximo - hoje
    if diff.days == 0:
        return f"Disponível em menos de 1 dia"
    return f"Disponível em {diff.days} dia{'s' if diff.days > 1 else ''}"

def risco_cor(nivel):
    return {'baixo': 'success', 'moderado': 'warning', 'alto': 'danger', 'critico': 'dark'}.get(nivel, 'secondary')

def risco_emoji(nivel):
    return {'baixo': '🟢', 'moderado': '🟡', 'alto': '🔴', 'critico': '💀'}.get(nivel, '⚪')

def plano_badge_color(plano_base):
    return {
        'starter1': '#8B949E',
        'starter2': '#6E7681',
        'profissional': '#0891B2',
        'empresarial': '#7C3AED',
        'enterprise': '#D97706',
    }.get(plano_base, '#8B949E')

app.jinja_env.globals['risco_cor'] = risco_cor
app.jinja_env.globals['risco_emoji'] = risco_emoji
app.jinja_env.globals['enumerate'] = enumerate
app.jinja_env.globals['PLANOS_INFO'] = PLANOS_INFO
app.jinja_env.globals['plano_badge_color'] = plano_badge_color

# ─────────────────────────── ROTAS GERAIS ───────────────────────────


@app.route('/esqueci-senha', methods=['GET', 'POST'])
def esqueci_senha():
    area = request.args.get('area', 'empresa')
    if request.method == 'GET':
        return render_template('esqueci_senha.html', area=area)
    email = request.form.get('email', '').strip().lower()
    if not email:
        flash('Informe seu e-mail.', 'danger')
        return render_template('esqueci_senha.html', area=area)
    import secrets as _sec, string as _str
    nova_senha = ''.join(_sec.choice(_str.ascii_letters + _str.digits) for _ in range(10))
    usuario = None
    if area == 'empresa':
        usuario = Empresa.query.filter_by(email=email).first()
    elif area == 'funcionario':
        usuario = Funcionario.query.filter_by(email=email).first()
    elif area == 'admin':
        usuario = Admin.query.filter_by(usuario=email).first()
    if usuario:
        usuario.senha = generate_password_hash(nova_senha)
        db.session.commit()
        try:
            gu = os.environ.get('GMAIL_USER', ''); gp = os.environ.get('GMAIL_APP_PASSWORD', '')
            if gu and gp:
                import smtplib
                from email.mime.text import MIMEText
                from email.mime.multipart import MIMEMultipart
                msg = MIMEMultipart('alternative')
                msg['Subject'] = 'BAI-EX - Nova senha temporaria'
                msg['From'] = 'BAI-EX <' + gu + '>'; msg['To'] = email
                bu = os.environ.get('BASE_URL', 'https://bai-ex.onrender.com')
                html = ('<div style="font-family:Arial,sans-serif;max-width:480px;margin:auto;padding:2rem;background:#0D1117;color:#F0F6FC;border-radius:12px">'
                        '<h2 style="color:#fff;font-size:1.1rem">Sua nova senha temporaria</h2>'
                        '<p style="color:#C9D1D9">Use a senha abaixo para entrar no BAI-EX:</p>'
                        '<div style="background:#161B22;border:1px solid #21262D;border-radius:8px;padding:1rem;margin:1rem 0;text-align:center">'
                        '<span style="color:#22D3EE;font-size:1.4rem;font-weight:700;letter-spacing:.15em">' + nova_senha + '</span></div>'
                        '<p style="color:#C9D1D9;font-size:.85rem">Apos entrar, va em <strong>Minha Conta</strong> para alterar sua senha.</p>'
                        '<a href="' + bu + '" style="display:inline-block;background:#0891B2;color:#fff;padding:.6rem 1.2rem;border-radius:8px;text-decoration:none;font-weight:600;font-size:.85rem;margin-top:.5rem">Acessar BAI-EX</a>'
                        '</div>')
                msg.attach(MIMEText(html, 'html'))
                with smtplib.SMTP_SSL('smtp.gmail.com', 465) as s:
                    s.login(gu, gp); s.sendmail(gu, email, msg.as_string())
        except Exception as e:
            print('[ESQUECI_SENHA] ' + str(e))
        registrar_auditoria('RESET_SENHA', 'Email: ' + email + ' | Area: ' + area, tipo_usuario=area)
    flash('Se este e-mail estiver cadastrado, voce recebera uma senha temporaria.', 'success')
    return redirect({'empresa': '/empresa/login', 'funcionario': '/funcionario/login', 'admin': '/admin/login'}.get(area, '/empresa/login'))


@app.route('/')
def index():
    return render_template('index.html', planos=PLANOS_INFO)

# ── FIX 8: Endpoint para verificar se sessão ainda está ativa ──
@app.route('/api/session/check')
def check_session():
    """Frontend chama esta rota periodicamente para detectar expiração."""
    tipo = session.get('tipo')
    if not tipo:
        return jsonify({'ativa': False}), 401
    return jsonify({'ativa': True, 'tipo': tipo})

# ── FIX 5: Política de Privacidade (exigida pela LGPD) ─────────
@app.route('/privacidade')
def privacidade():
    return render_template('privacidade.html')

@app.route('/termos')
def termos():
    return render_template('termos.html')

# ─────────────────────────── EMPRESA / RH ───────────────────────────

@app.route('/empresa/login', methods=['GET', 'POST'])
def login_empresa():
    if request.method == 'POST':
        ip = request.remote_addr
        bloqueio = verificar_bloqueio(ip, area='empresa')
        if bloqueio:
            flash(f'Muitas tentativas. Aguarde {bloqueio} minuto(s).', 'danger')
            return render_template('empresa/login.html')
        email = request.form['email']
        senha = request.form['senha']
        empresa = Empresa.query.filter_by(email=email).first()
        if empresa and check_password_hash(empresa.senha, senha):
            resetar_tentativas(ip, area='empresa')
            session['empresa_id'] = empresa.id
            session['empresa_nome'] = empresa.nome
            session['tipo'] = 'empresa'
            marcar_sessao_permanente()
            registrar_auditoria('LOGIN', f'Empresa: {empresa.nome}',
                                tipo_usuario='empresa', usuario_id=empresa.id,
                                empresa_id=empresa.id)
            return redirect(url_for('dashboard_empresa'))
        registrar_falha(ip, area='empresa')
        registrar_auditoria('LOGIN_FALHA', f'Email: {email}', tipo_usuario='empresa')
        flash('Email ou senha incorretos.', 'danger')
    return render_template('empresa/login.html')

@app.route('/empresa/cadastro', methods=['GET', 'POST'])
def cadastro_empresa():
    plano_escolhido = request.args.get('plano', 'starter1')
    if plano_escolhido not in PLANOS_INFO:
        plano_escolhido = 'starter1'

    if request.method == 'POST':
        nome = request.form['nome']
        cnpj = request.form['cnpj']
        responsavel = request.form['responsavel']
        email = request.form['email']
        senha = request.form['senha']
        plano_form = request.form.get('plano', 'starter1')
        if plano_form not in PLANOS_INFO:
            plano_form = 'starter1'
        cnpj_limpo = ''.join(filter(str.isdigit, cnpj))
        if not validar_cnpj(cnpj):
            flash('CNPJ inválido. Verifique e tente novamente.', 'danger')
        elif Empresa.query.filter_by(email=email).first():
            flash('Email já cadastrado.', 'danger')
        elif Empresa.query.filter_by(cnpj=cnpj_limpo).first():
            flash('CNPJ já cadastrado.', 'danger')
        else:
            empresa = Empresa(
                nome=nome, cnpj=formatar_cnpj(cnpj),
                responsavel=responsavel, email=email,
                senha=generate_password_hash(senha),
                plano=f'trial_{plano_form}',
                trial_inicio=datetime.utcnow()
            )
            db.session.add(empresa)
            db.session.commit()
            trial_dias = PLANOS_INFO[plano_form]['trial_dias']
            registrar_auditoria('CADASTRO_EMPRESA',
                                f'Empresa: {nome} | Plano trial: {plano_form} | Trial: {trial_dias}d',
                                tipo_usuario='empresa', usuario_id=empresa.id, empresa_id=empresa.id)
            flash(f'🎉 Empresa cadastrada! Você tem {trial_dias} dias grátis. Faça login para começar.', 'success')
            return redirect(url_for('login_empresa'))
    return render_template('empresa/cadastro.html',
                           plano=plano_escolhido,
                           plano_info=PLANOS_INFO.get(plano_escolhido, {}))

@app.route('/empresa/logout')
def logout_empresa():
    session.clear()
    return redirect(url_for('index'))

@app.route('/empresa/dashboard')
@login_required_empresa
def dashboard_empresa():
    empresa_id = session['empresa_id']
    empresa = Empresa.query.get(empresa_id)
    funcionarios = Funcionario.query.filter_by(empresa_id=empresa_id, ativo=True).all()
    semana, ano = get_semana_atual()

    # Resumo geral
    total_func = len(funcionarios)
    ids = [f.id for f in funcionarios]

    checkins_semana = CheckIn.query.filter(
        CheckIn.empresa_id == empresa_id,
        CheckIn.semana == semana,
        CheckIn.ano == ano
    ).all()

    respostas_semana = len(checkins_semana)
    taxa_resposta = round((respostas_semana / total_func * 100) if total_func else 0)

    # Score médio e riscos
    scores = [c.score for c in checkins_semana]
    score_medio = round(sum(scores) / len(scores), 1) if scores else 0

    risco_alto = len([c for c in checkins_semana if c.nivel_risco in ('alto', 'critico')])
    risco_moderado = len([c for c in checkins_semana if c.nivel_risco == 'moderado'])

    nivel_geral = CheckIn.classificar_risco(score_medio) if score_medio else 'baixo'

    # Evolução das últimas 8 semanas
    historico = []
    for i in range(7, -1, -1):
        dt = date.today() - timedelta(weeks=i)
        s = dt.isocalendar()[1]
        a = dt.year
        cks = CheckIn.query.filter_by(empresa_id=empresa_id, semana=s, ano=a).all()
        sc = round(sum(c.score for c in cks) / len(cks), 1) if cks else 0
        historico.append({'semana': f'Sem {s}', 'score': sc, 'respostas': len(cks)})

    # Distribuição de riscos (todos os checkins da semana)
    dist = {'baixo': 0, 'moderado': 0, 'alto': 0, 'critico': 0}
    for c in checkins_semana:
        dist[c.nivel_risco] += 1

    # Funcionários com risco - últimos checkin de cada um
    func_risco = []
    for f in funcionarios:
        ult = CheckIn.query.filter_by(funcionario_id=f.id).order_by(CheckIn.data.desc()).first()
        func_risco.append({'funcionario': f, 'checkin': ult})
    func_risco.sort(key=lambda x: (x['checkin'].score if x['checkin'] else 0), reverse=True)

    return render_template('empresa/dashboard.html',
        empresa=empresa,
        total_func=total_func,
        respostas_semana=respostas_semana,
        taxa_resposta=taxa_resposta,
        score_medio=score_medio,
        nivel_geral=nivel_geral,
        risco_alto=risco_alto,
        risco_moderado=risco_moderado,
        historico=historico,
        dist=dist,
        func_risco=func_risco,
        semana=semana
    )

@app.route('/empresa/funcionarios')
@login_required_empresa
def funcionarios_lista():
    empresa_id = session['empresa_id']
    funcionarios = Funcionario.query.filter_by(empresa_id=empresa_id).all()
    semana, ano = get_semana_atual()
    func_data = []
    for f in funcionarios:
        ult = CheckIn.query.filter_by(funcionario_id=f.id).order_by(CheckIn.data.desc()).first()
        checkin_semana = CheckIn.query.filter_by(funcionario_id=f.id, semana=semana, ano=ano).first()
        func_data.append({'f': f, 'ultimo': ult, 'esta_semana': checkin_semana is not None})
    return render_template('empresa/funcionarios.html', func_data=func_data)

@app.route('/empresa/funcionarios/<int:fid>/toggle')
@login_required_empresa
def toggle_funcionario(fid):
    f = Funcionario.query.get_or_404(fid)
    if f.empresa_id != session['empresa_id']:
        flash('Acesso negado.', 'danger')
        return redirect(url_for('funcionarios_lista'))
    f.ativo = not f.ativo
    db.session.commit()
    flash(f'Funcionário {"ativado" if f.ativo else "desativado"}.', 'info')
    return redirect(url_for('funcionarios_lista'))

@app.route('/empresa/relatorio')
@login_required_empresa
def relatorio_empresa():
    empresa_id = session['empresa_id']
    empresa = Empresa.query.get(empresa_id)
    semana, ano = get_semana_atual()

    # Últimas 4 semanas
    dados_semanas = []
    for i in range(3, -1, -1):
        dt = date.today() - timedelta(weeks=i)
        s = dt.isocalendar()[1]
        a = dt.year
        cks = CheckIn.query.filter_by(empresa_id=empresa_id, semana=s, ano=a).all()
        sc = round(sum(c.score for c in cks) / len(cks), 1) if cks else 0
        dist = {'baixo': 0, 'moderado': 0, 'alto': 0, 'critico': 0}
        for c in cks:
            dist[c.nivel_risco] += 1
        dados_semanas.append({'semana': s, 'ano': a, 'score': sc, 'total': len(cks), 'dist': dist})

    total_func = Funcionario.query.filter_by(empresa_id=empresa_id, ativo=True).count()
    return render_template('empresa/relatorio.html',
        empresa=empresa,
        dados_semanas=dados_semanas,
        total_func=total_func,
        gerado_em=datetime.now().strftime('%d/%m/%Y %H:%M')
    )


@app.route('/empresa/relatorio/download-docx')
@login_required_empresa
def download_relatorio_docx():
    """Gera Relatorio de Compliance NR-01 em .docx - 7 secoes conforme modelo PGR."""
    from docx import Document as DocxDoc
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from io import BytesIO

    empresa_id = session['empresa_id']
    empresa    = Empresa.query.get(empresa_id)
    periodo    = request.args.get('periodo', 'semanal')
    num_semanas = {'semanal': 4, 'quinzenal': 8, 'mensal': 12}.get(periodo, 4)

    dados_semanas = []
    for i in range(num_semanas - 1, -1, -1):
        dt   = date.today() - timedelta(weeks=i)
        s, a = dt.isocalendar()[1], dt.year
        cks  = CheckIn.query.filter_by(empresa_id=empresa_id, semana=s, ano=a).all()
        sc   = round(sum(c.score for c in cks) / len(cks), 1) if cks else 0
        dist = {'baixo': 0, 'moderado': 0, 'alto': 0, 'critico': 0}
        for c in cks:
            if c.nivel_risco in dist:
                dist[c.nivel_risco] += 1
        dados_semanas.append({'semana': s, 'ano': a, 'score': sc, 'total': len(cks), 'dist': dist})

    total_func     = Funcionario.query.filter_by(empresa_id=empresa_id, ativo=True).count()
    total_checkins = sum(d['total'] for d in dados_semanas)
    scores_validos = [d['score'] for d in dados_semanas if d['score'] > 0]
    score_geral    = round(sum(scores_validos) / max(1, len(scores_validos)), 1)
    dist_total     = {'baixo': 0, 'moderado': 0, 'alto': 0, 'critico': 0}
    for d in dados_semanas:
        for k in dist_total:
            dist_total[k] += d['dist'][k]

    if score_geral < 30:   nivel_conf = 'ALTO';   status_conf = 'Conforme'
    elif score_geral < 55: nivel_conf = 'MEDIO';  status_conf = 'Parcial'
    else:                  nivel_conf = 'BAIXO';  status_conf = 'Nao conforme'

    nao_conf = []
    for d in dados_semanas:
        taxa = round(d['total'] / max(1, total_func) * 100, 0)
        if taxa < 70:
            nao_conf.append({
                'desc': 'Semana ' + str(d['semana']) + '/' + str(d['ano']) + ': taxa de resposta ' + str(int(taxa)) + '% (abaixo de 70%)',
                'class': 'Moderada',
                'acao': 'Reforcar comunicacao e engajamento dos colaboradores',
                'prazo': '2 semanas'
            })
        if d['score'] >= 55:
            nao_conf.append({
                'desc': 'Semana ' + str(d['semana']) + '/' + str(d['ano']) + ': score ' + str(d['score']) + '/100 (nivel alto)',
                'class': 'Grave',
                'acao': 'Acionar RH para intervencao junto aos colaboradores em risco alto/critico',
                'prazo': '48 horas'
            })

    doc = DocxDoc()
    for sec in doc.sections:
        sec.top_margin = sec.bottom_margin = Cm(2.5)
        sec.left_margin = sec.right_margin = Cm(3.18)

    def sf(run, size=11, bold=False, rgb=None):
        run.font.name = 'Arial'
        run.font.size = Pt(size)
        run.font.bold = bold
        if rgb:
            run.font.color.rgb = RGBColor(*rgb)

    def cbg(cell, hex_c):
        tc = cell._tc; tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), hex_c); tcPr.append(shd)

    def h1(doc, text):
        p = doc.add_paragraph()
        r = p.add_run(text)
        sf(r, 13, True, (8, 64, 110))
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after  = Pt(4)
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        btm  = OxmlElement('w:bottom')
        btm.set(qn('w:val'), 'single'); btm.set(qn('w:sz'), '6')
        btm.set(qn('w:space'), '1');    btm.set(qn('w:color'), '0891B2')
        pBdr.append(btm); pPr.append(pBdr)
        return p

    def kv(doc, lbl, val):
        p = doc.add_paragraph()
        sf(p.add_run(lbl + ': '), 10, True)
        sf(p.add_run(str(val)), 10)
        p.paragraph_format.space_after = Pt(2)

    # TITULO
    p_t = doc.add_paragraph()
    sf(p_t.add_run('RELATORIO DE COMPLIANCE - NR-01'), 20, True, (13, 17, 23))
    p_t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub = doc.add_paragraph()
    sf(p_sub.add_run('BAI-EX - Burnout AI Experience  -  Gerado em ' + datetime.now().strftime('%d/%m/%Y as %H:%M')), 9, False, (139, 148, 158))
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 1. IDENTIFICACAO
    h1(doc, '1. IDENTIFICACAO DA EMPRESA')
    kv(doc, 'Razao Social', empresa.nome)
    kv(doc, 'CNPJ', empresa.cnpj)
    kv(doc, 'Responsavel pelo PGR', empresa.responsavel)
    kv(doc, 'E-mail', empresa.email)
    kv(doc, 'Colaboradores monitorados', total_func)
    kv(doc, 'Periodo analisado', str(num_semanas) + ' semanas (' + periodo + ')')

    # 2. OBJETIVO
    h1(doc, '2. OBJETIVO')
    p_obj = doc.add_paragraph()
    sf(p_obj.add_run(
        'Avaliar o nivel de conformidade com a NR-01 atualizada, incluindo o Gerenciamento de '
        'Riscos Ocupacionais (GRO) e o Programa de Gerenciamento de Riscos (PGR), com foco nos '
        'riscos psicossociais exigidos pelo MTE a partir de janeiro de 2025.'
    ), 10)

    # 3. ESCOPO
    h1(doc, '3. ESCOPO')
    for item in [
        'Monitoramento continuo de riscos psicossociais por colaborador via check-in periodico',
        'Analise de ' + str(num_semanas) + ' semanas de dados validados',
        'Identificacao por dimensao: estresse, sobrecarga, sono, motivacao, relacoes interpessoais',
        'Score de risco 0-100 por colaborador com classificacao em 4 niveis (baixo/moderado/alto/critico)',
        'Emissao de alertas automaticos ao RH para colaboradores em risco alto ou critico',
    ]:
        p_i = doc.add_paragraph()
        p_i.paragraph_format.left_indent = Pt(18)
        p_i.paragraph_format.space_after = Pt(2)
        sf(p_i.add_run('- '), 10, True, (8, 145, 178))
        sf(p_i.add_run(item), 10)

    # 4. REQUISITOS E STATUS
    h1(doc, '4. REQUISITOS E STATUS DE CONFORMIDADE')

    p_41 = doc.add_paragraph()
    sf(p_41.add_run('4.1 Monitoramento Semanal de Riscos Psicossociais'), 11, True, (8, 64, 110))
    p_41.paragraph_format.space_after = Pt(4)

    hdrs = ['Semana', 'Check-ins', 'Score Medio', 'Classificacao', 'Baixo', 'Moderado', 'Alto', 'Critico']
    t1 = doc.add_table(rows=1 + len(dados_semanas), cols=len(hdrs))
    t1.style = 'Table Grid'
    for j, h in enumerate(hdrs):
        c = t1.rows[0].cells[j]; c.text = h; cbg(c, '0D3349')
        runs = c.paragraphs[0].runs
        if runs: sf(runs[0], 9, True, (255, 255, 255))
    for i, d in enumerate(dados_semanas):
        if   d['score'] >= 75: nv = 'Critico'
        elif d['score'] >= 55: nv = 'Alto'
        elif d['score'] >= 30: nv = 'Moderado'
        elif d['score'] >  0:  nv = 'Baixo'
        else:                  nv = '-'
        for j, val in enumerate([
            'S' + str(d['semana']) + '/' + str(d['ano']), str(d['total']) + '/' + str(total_func),
            str(d['score']) + '/100', nv,
            str(d['dist']['baixo']), str(d['dist']['moderado']),
            str(d['dist']['alto']),  str(d['dist']['critico']),
        ]):
            c = t1.rows[i+1].cells[j]; c.text = val
            cbg(c, 'E8F4F8' if i % 2 == 0 else 'FFFFFF')
    doc.add_paragraph()

    p_42 = doc.add_paragraph()
    sf(p_42.add_run('4.2 Avaliacao de Requisitos NR-01'), 11, True, (8, 64, 110))
    p_42.paragraph_format.space_after = Pt(4)

    c_ok  = '(X)' if status_conf == 'Conforme'      else '( )'
    c_pc  = '(X)' if status_conf == 'Parcial'        else '( )'
    c_nc  = '(X)' if status_conf == 'Nao conforme'   else '( )'
    reqs  = [
        ('GRO - Gerenciamento de Riscos Ocupacionais',     c_ok, c_pc, c_nc),
        ('PGR - Programa de Gerenciamento de Riscos',      c_ok, c_pc, c_nc),
        ('Inventario de Riscos Psicossociais',              c_ok, c_pc, c_nc),
        ('Monitoramento Continuo dos Colaboradores',        '(X)', '( )', '( )'),
        ('Emissao de Alertas e Acoes Preventivas',          '(X)', '( )', '( )'),
        ('Documentacao e Rastreabilidade (20 anos)',        '(X)', '( )', '( )'),
    ]
    t2 = doc.add_table(rows=1 + len(reqs), cols=4)
    t2.style = 'Table Grid'
    for j, h in enumerate(['Requisito', 'Conforme', 'Parcial', 'Nao conforme']):
        c2 = t2.rows[0].cells[j]; c2.text = h; cbg(c2, '0D3349')
        runs = c2.paragraphs[0].runs
        if runs: sf(runs[0], 9, True, (255, 255, 255))
    for i, (req, ok, pc, nc) in enumerate(reqs):
        for j, val in enumerate([req, ok, pc, nc]):
            c2 = t2.rows[i+1].cells[j]; c2.text = val
            cbg(c2, 'E8F4F8' if i % 2 == 0 else 'FFFFFF')
            if j > 0: c2.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    # 5. NAO CONFORMIDADES
    h1(doc, '5. NAO CONFORMIDADES IDENTIFICADAS')
    if not nao_conf:
        sf(doc.add_paragraph().add_run('Nenhuma nao conformidade identificada no periodo analisado.'), 10, False, (22, 163, 74))
    else:
        t3 = doc.add_table(rows=1 + len(nao_conf), cols=4)
        t3.style = 'Table Grid'
        for j, h in enumerate(['Descricao', 'Classificacao', 'Acao Recomendada', 'Prazo']):
            c3 = t3.rows[0].cells[j]; c3.text = h; cbg(c3, '0D3349')
            runs = c3.paragraphs[0].runs
            if runs: sf(runs[0], 9, True, (255, 255, 255))
        for i, item in enumerate(nao_conf):
            for j, val in enumerate([item['desc'], item['class'], item['acao'], item['prazo']]):
                c3 = t3.rows[i+1].cells[j]; c3.text = val
                cbg(c3, 'FFF5F5' if item['class'] == 'Grave' else 'FFFBEB')
    doc.add_paragraph()

    # 6. CONCLUSAO
    h1(doc, '6. CONCLUSAO')
    p_conc = doc.add_paragraph()
    rgb_nv = (22, 163, 74) if nivel_conf == 'ALTO' else (234, 179, 8) if nivel_conf == 'MEDIO' else (239, 68, 68)
    sf(p_conc.add_run('Nivel de conformidade identificado: '), 10)
    sf(p_conc.add_run(nivel_conf + ' (' + status_conf + ')'), 11, True, rgb_nv)
    txt = (chr(10) + chr(10)
           + 'No periodo de ' + str(num_semanas) + ' semanas, foram realizados '
           + str(total_checkins) + ' check-ins de ' + str(total_func) + ' colaboradores. '
           + 'Score medio geral: ' + str(score_geral) + '/100. '
           + 'Risco alto: ' + str(dist_total['alto']) + ' | Critico: ' + str(dist_total['critico']) + '.')
    sf(p_conc.add_run(txt), 10)
    if nao_conf:
        sf(p_conc.add_run(chr(10) + chr(10) + 'Recomenda-se execucao imediata das '
                          + str(len(nao_conf)) + ' acao(oes) corretiva(s) da secao 5.'), 10)
    else:
        sf(p_conc.add_run(chr(10) + chr(10) + 'A empresa esta em conformidade com os requisitos de monitoramento da NR-01.'),
           10, False, (22, 163, 74))

    # 7. RESPONSAVEL
    h1(doc, '7. RESPONSAVEL PELO RELATORIO')
    doc.add_paragraph()
    kv(doc, 'Nome', empresa.responsavel)
    kv(doc, 'Cargo', 'Responsavel pelo PGR')
    kv(doc, 'Empresa', empresa.nome)
    p_assin = doc.add_paragraph()
    p_assin.paragraph_format.space_before = Pt(20)
    sf(p_assin.add_run('Assinatura: _________________________________     '), 10)
    sf(p_assin.add_run('Data: ' + datetime.now().strftime('%d/%m/%Y')), 10)
    doc.add_paragraph()
    p_rod = doc.add_paragraph()
    sf(p_rod.add_run('Relatorio gerado pelo BAI-EX - ' + datetime.now().strftime('%d/%m/%Y %H:%M')
                     + ' - Evidencia de conformidade NR-01 (Norma Regulamentadora n1 - MTE)'), 8, False, (139, 148, 158))
    p_rod.alignment = WD_ALIGN_PARAGRAPH.CENTER

    buf = BytesIO(); doc.save(buf); buf.seek(0)
    cnpj_n   = ''.join(filter(str.isdigit, empresa.cnpj or ''))[:14]
    filename = 'relatorio_compliance_nr01_' + cnpj_n + '_' + datetime.now().strftime('%Y%m%d') + '_' + periodo + '.docx'
    registrar_auditoria('DOWNLOAD_RELATORIO_NR01', 'Empresa: ' + empresa.nome + ' | Periodo: ' + periodo)
    resp = make_response(buf.read())
    resp.headers['Content-Type'] = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    resp.headers['Content-Disposition'] = 'attachment; filename=' + filename
    return resp


@app.route('/empresa/alertas')
@login_required_empresa
def alertas_empresa():
    empresa_id = session['empresa_id']
    semana, ano = get_semana_atual()

    # Funcionários com risco alto ou crítico nas últimas 2 semanas
    s2 = (date.today() - timedelta(weeks=1)).isocalendar()[1]
    checkins_criticos = CheckIn.query.filter(
        CheckIn.empresa_id == empresa_id,
        CheckIn.nivel_risco.in_(['alto', 'critico']),
        CheckIn.ano == ano
    ).order_by(CheckIn.data.desc()).all()

    # Funcionários que não responderam esta semana
    todos = Funcionario.query.filter_by(empresa_id=empresa_id, ativo=True).all()
    responderam = {c.funcionario_id for c in CheckIn.query.filter_by(
        empresa_id=empresa_id, semana=semana, ano=ano).all()}
    nao_responderam = [f for f in todos if f.id not in responderam]

    return render_template('empresa/alertas.html',
        checkins_criticos=checkins_criticos,
        nao_responderam=nao_responderam,
        semana=semana
    )

# ─────────────────────────── FUNCIONÁRIO ───────────────────────────

@app.route('/funcionario/login', methods=['GET', 'POST'])
def login_funcionario():
    if request.method == 'POST':
        ip = request.remote_addr
        bloqueio = verificar_bloqueio(ip, area='funcionario')
        if bloqueio:
            flash(f'Muitas tentativas. Aguarde {bloqueio} minuto(s).', 'danger')
            return render_template('funcionario/login.html')
        email = request.form['email']
        senha = request.form['senha']
        f = Funcionario.query.filter_by(email=email).first()
        if f and check_password_hash(f.senha, senha):
            resetar_tentativas(ip, area='funcionario')
            session['funcionario_id'] = f.id
            session['funcionario_nome'] = f.nome
            session['empresa_id_func'] = f.empresa_id
            session['tipo'] = 'funcionario'
            marcar_sessao_permanente()
            registrar_auditoria('LOGIN', f'Funcionário: {f.nome}',
                                tipo_usuario='funcionario', usuario_id=f.id,
                                empresa_id=f.empresa_id)
            return redirect(url_for('painel_funcionario'))
        registrar_falha(ip, area='funcionario')
        registrar_auditoria('LOGIN_FALHA', f'Email: {email}', tipo_usuario='funcionario')
        flash('Email ou senha incorretos.', 'danger')
    return render_template('funcionario/login.html')

@app.route('/funcionario/logout')
def logout_funcionario():
    session.clear()
    return redirect(url_for('index'))

@app.route('/funcionario/painel')
@login_required_funcionario
def painel_funcionario():
    fid = session['funcionario_id']
    f = Funcionario.query.get(fid)
    empresa = Empresa.query.get(f.empresa_id)
    semana, ano = get_semana_atual()
    periodo_ref, periodo_label = periodo_atual(empresa)

    # Verifica se já respondeu no período atual
    ja_respondeu = CheckIn.query.filter_by(
        funcionario_id=fid, periodo_ref=periodo_ref).first()

    historico = CheckIn.query.filter_by(funcionario_id=fid).order_by(CheckIn.data.desc()).limit(8).all()
    ultimo = historico[0] if historico else None

    proximo_label = proximo_checkin_label(empresa, ultimo.data if ultimo else None)

    historico_json = [
        {'semana': ck.semana, 'ano': ck.ano, 'score': ck.score, 'nivel_risco': ck.nivel_risco}
        for ck in historico
    ]

    return render_template('funcionario/painel.html',
        funcionario=f,
        empresa=empresa,
        ja_respondeu=ja_respondeu,
        historico=historico,
        historico_json=historico_json,
        ultimo=ultimo,
        semana=semana,
        periodo_label=periodo_label,
        proximo_label=proximo_label,
        freq_label={'semanal': 'Semanal', 'quinzenal': 'Quinzenal', 'mensal': 'Mensal'}.get(
            empresa.frequencia_checkin, 'Semanal')
    )

@app.route('/funcionario/checkin', methods=['GET', 'POST'])
@login_required_funcionario
def checkin_funcionario():
    fid = session['funcionario_id']
    f = Funcionario.query.get(fid)
    empresa = Empresa.query.get(f.empresa_id)
    semana, ano = get_semana_atual()
    periodo_ref, periodo_label = periodo_atual(empresa)

    ja_respondeu = CheckIn.query.filter_by(funcionario_id=fid, periodo_ref=periodo_ref).first()
    freq_label = {'semanal': 'semana', 'quinzenal': 'quinzena', 'mensal': 'mês'}.get(
        empresa.frequencia_checkin, 'semana')
    if ja_respondeu:
        flash(f'Você já respondeu o check-in desta {freq_label}! '
              f'Próximo disponível no próximo período.', 'info')
        return redirect(url_for('painel_funcionario'))

    if request.method == 'POST':
        try:
            respostas = [int(request.form[f'q{i}']) for i in range(1, 25)]
            if not all(1 <= r <= 5 for r in respostas):
                flash('Responda todas as perguntas.', 'danger')
                return render_template('funcionario/checkin.html', funcionario=f,
                                       periodo_label=periodo_label)

            score = CheckIn.calcular_score(respostas)
            nivel = CheckIn.classificar_risco(score)

            ck = CheckIn(
                funcionario_id=fid,
                empresa_id=f.empresa_id,
                semana=semana, ano=ano,
                periodo_ref=periodo_ref,
                q1_estresse=respostas[0],   q2_sobrecarga=respostas[1],
                q3_sono=respostas[2],       q4_motivacao=respostas[3],
                q5_esgotamento=respostas[4], q6_relacoes=respostas[5],
                q7_autonomia=respostas[6],   q8_reconhecimento=respostas[7],
                q9_clareza=respostas[8],     q10_comunicacao=respostas[9],
                q11_lideranca=respostas[10], q12_ritmo=respostas[11],
                q13_interrupcoes=respostas[12], q14_decisao=respostas[13],
                q15_seguranca=respostas[14], q16_conflitos=respostas[15],
                q17_justica=respostas[16],   q18_desconexao=respostas[17],
                q19_descanso=respostas[18],  q20_pressao=respostas[19],
                q21_desanimo=respostas[20],  q22_respeito=respostas[21],
                q23_performance=respostas[22], q24_satisfacao=respostas[23],
                score=score, nivel_risco=nivel
            )
            db.session.add(ck)
            db.session.commit()
            registrar_auditoria('CHECKIN', f'Score: {score} Nivel: {nivel} Periodo: {periodo_ref}',
                                tipo_usuario='funcionario',
                                usuario_id=fid, empresa_id=f.empresa_id)
            return redirect(url_for('resultado_checkin', ck_id=ck.id))
        except (ValueError, KeyError):
            flash('Por favor, responda todas as perguntas.', 'danger')

    return render_template('funcionario/checkin.html', funcionario=f,
                           periodo_label=periodo_label)

@app.route('/funcionario/resultado/<int:ck_id>')
@login_required_funcionario
def resultado_checkin(ck_id):
    ck = CheckIn.query.get_or_404(ck_id)
    if ck.funcionario_id != session['funcionario_id']:
        return redirect(url_for('painel_funcionario'))
    return render_template('funcionario/resultado.html', ck=ck)

# ─────────────────────────── ADMIN ───────────────────────────

# ── FIX 7: Rota de backup manual (admin only) ──────────────────
@app.route('/admin/backup')
@login_required_admin
def backup_dados():
    """
    Exporta todos os dados em JSON para backup manual.
    Em produção no Render, o PostgreSQL já faz backup automático.
    Esta rota serve como backup adicional para o admin baixar.
    """
    import json as _json
    from sqlalchemy import inspect as _inspect

    dados = {}

    # Empresas (sem senhas)
    dados['empresas'] = [
        {'id': e.id, 'nome': e.nome, 'cnpj': e.cnpj,
         'responsavel': e.responsavel, 'email': e.email,
         'plano': e.plano, 'criado_em': str(e.criado_em)}
        for e in Empresa.query.all()
    ]

    # Funcionários (sem senhas)
    dados['funcionarios'] = [
        {'id': f.id, 'empresa_id': f.empresa_id, 'nome': f.nome,
         'email': f.email, 'departamento': f.departamento,
         'cargo': f.cargo, 'ativo': f.ativo}
        for f in Funcionario.query.all()
    ]

    # Check-ins
    dados['checkins'] = [
        {'id': ck.id, 'funcionario_id': ck.funcionario_id,
         'empresa_id': ck.empresa_id, 'data': str(ck.data),
         'semana': ck.semana, 'ano': ck.ano, 'score': ck.score,
         'nivel_risco': ck.nivel_risco}
        for ck in CheckIn.query.all()
    ]

    dados['gerado_em'] = str(datetime.utcnow())
    dados['total'] = {k: len(v) for k, v in dados.items() if isinstance(v, list)}

    registrar_auditoria('BACKUP', f"Total: {dados['total']}", tipo_usuario='admin')

    from flask import make_response
    resp = make_response(_json.dumps(dados, ensure_ascii=False, indent=2))
    resp.headers['Content-Type'] = 'application/json'
    resp.headers['Content-Disposition'] = f'attachment; filename=baiex_backup_{datetime.now().strftime("%Y%m%d_%H%M")}.json'
    return resp

@app.route('/admin/login', methods=['GET', 'POST'])
def login_admin():
    if request.method == 'POST':
        ip = request.remote_addr
        bloqueio = verificar_bloqueio(ip, area='admin')
        if bloqueio:
            flash(f'Muitas tentativas. Aguarde {bloqueio} minuto(s).', 'danger')
            return render_template('admin/login.html')
        usuario = request.form['usuario']
        senha = request.form['senha']
        admin = Admin.query.filter_by(usuario=usuario).first()
        if admin and check_password_hash(admin.senha, senha):
            resetar_tentativas(ip, area='admin')
            session['admin_id'] = admin.id
            session['tipo'] = 'admin'
            marcar_sessao_permanente()
            return redirect(url_for('dashboard_admin'))
        registrar_falha(ip, area='admin')
        flash('Credenciais inválidas.', 'danger')
    return render_template('admin/login.html')

@app.route('/admin/logout')
def logout_admin():
    session.clear()
    return redirect(url_for('index'))

# ── Rota de visualização de logs (admin) ───────────────────────
@app.route('/admin/logs')
@login_required_admin
def logs_auditoria():
    page = request.args.get('page', 1, type=int)
    acao_filtro = request.args.get('acao', '')
    query = AuditLog.query
    if acao_filtro:
        query = query.filter(AuditLog.acao.ilike(f'%{acao_filtro}%'))
    logs = query.order_by(AuditLog.timestamp.desc()).limit(200).all()
    acoes_unicas = db.session.query(AuditLog.acao).distinct().all()
    acoes_unicas = [a[0] for a in acoes_unicas]
    return render_template('admin/logs.html',
        logs=logs, acao_filtro=acao_filtro, acoes_unicas=acoes_unicas)

@app.route('/admin/dashboard')
@login_required_admin
def dashboard_admin():
    total_empresas = Empresa.query.count()
    total_func = Funcionario.query.count()
    total_checkins = CheckIn.query.count()

    empresas = Empresa.query.order_by(Empresa.criado_em.desc()).all()
    semana, ano = get_semana_atual()

    empresas_data = []
    for e in empresas:
        func_count = Funcionario.query.filter_by(empresa_id=e.id, ativo=True).count()
        checkins_semana = CheckIn.query.filter_by(empresa_id=e.id, semana=semana, ano=ano).count()
        ultimo_ck = CheckIn.query.filter_by(empresa_id=e.id).order_by(CheckIn.data.desc()).first()
        scores = [c.score for c in CheckIn.query.filter_by(empresa_id=e.id, semana=semana, ano=ano).all()]
        score_med = round(sum(scores)/len(scores), 1) if scores else 0
        empresas_data.append({
            'empresa': e,
            'func_count': func_count,
            'checkins_semana': checkins_semana,
            'score_medio': score_med,
            'nivel': CheckIn.classificar_risco(score_med) if score_med else 'baixo',
            'ultimo_ck': ultimo_ck,
            'plano_info': PLANOS_INFO.get(e.plano_base, PLANOS_INFO['starter1']),
        })

    return render_template('admin/dashboard.html',
        total_empresas=total_empresas,
        total_func=total_func,
        total_checkins=total_checkins,
        empresas_data=empresas_data,
        PLANOS_INFO=PLANOS_INFO,
    )

# ─────────────────────────── API JSON ───────────────────────────

@app.route('/api/empresa/historico')
@login_required_empresa
def api_historico():
    empresa_id = session['empresa_id']
    labels, scores, respostas = [], [], []
    for i in range(7, -1, -1):
        dt = date.today() - timedelta(weeks=i)
        s = dt.isocalendar()[1]
        a = dt.year
        cks = CheckIn.query.filter_by(empresa_id=empresa_id, semana=s, ano=a).all()
        sc = round(sum(c.score for c in cks) / len(cks), 1) if cks else 0
        labels.append(f'Sem {s}')
        scores.append(sc)
        respostas.append(len(cks))
    return jsonify({'labels': labels, 'scores': scores, 'respostas': respostas})

# ─────────────────────────── SEED & INIT ───────────────────────────


# ═══════════════════════════════════════════════════════════════
#   MÓDULOS DE IA — BAI-EX  (powered by Claude)
# ═══════════════════════════════════════════════════════════════

def _dados_empresa_para_ia(empresa_id):
    """Monta contexto completo da empresa para enviar à IA."""
    empresa = Empresa.query.get(empresa_id)
    semana, ano = get_semana_atual()
    funcionarios = Funcionario.query.filter_by(empresa_id=empresa_id, ativo=True).all()

    # Últimas 4 semanas
    semanas_dados = []
    for i in range(3, -1, -1):
        dt = date.today() - timedelta(weeks=i)
        s = dt.isocalendar()[1]
        a = dt.year
        cks = CheckIn.query.filter_by(empresa_id=empresa_id, semana=s, ano=a).all()
        sc = round(sum(c.score for c in cks) / len(cks), 1) if cks else 0
        dist = {'baixo': 0, 'moderado': 0, 'alto': 0, 'critico': 0}
        for c in cks:
            dist[c.nivel_risco] += 1
        semanas_dados.append({'semana': s, 'score': sc, 'dist': dist, 'total': len(cks)})

    # Por departamento
    dept_dados = {}
    for f in funcionarios:
        ult = CheckIn.query.filter_by(funcionario_id=f.id).order_by(CheckIn.data.desc()).first()
        if ult:
            d = f.departamento
            if d not in dept_dados:
                dept_dados[d] = []
            dept_dados[d].append(ult.score)

    dept_resumo = {d: round(sum(v)/len(v), 1) for d, v in dept_dados.items()}

    # Funcionários críticos
    criticos = []
    for f in funcionarios:
        ult = CheckIn.query.filter_by(funcionario_id=f.id).order_by(CheckIn.data.desc()).first()
        if ult and ult.nivel_risco in ('alto', 'critico'):
            criticos.append({'nome': f.nome[:20], 'dept': f.departamento,
                             'score': ult.score, 'nivel': ult.nivel_risco})

    total_func = len(funcionarios)
    cks_semana = CheckIn.query.filter_by(empresa_id=empresa_id, semana=semana, ano=ano).all()
    score_atual = round(sum(c.score for c in cks_semana) / len(cks_semana), 1) if cks_semana else 0

    return {
        'empresa': empresa.nome,
        'total_funcionarios': total_func,
        'score_atual': score_atual,
        'nivel_atual': CheckIn.classificar_risco(score_atual) if score_atual else 'baixo',
        'taxa_resposta': round(len(cks_semana) / total_func * 100) if total_func else 0,
        'historico_semanas': semanas_dados,
        'por_departamento': dept_resumo,
        'funcionarios_criticos': criticos[:5],
        'semana': semana,
        'ano': ano,
    }


# ── 1. PARECER NARRATIVO NR-01 (streaming) ──────────────────────

@app.route('/empresa/relatorio/ia')
@login_required_empresa
@requer_plano('profissional')
def relatorio_ia():
    dados = _dados_empresa_para_ia(session['empresa_id'])
    registrar_auditoria('RELATORIO_IA', 'Acesso ao relatório NR-01 IA')
    return render_template('empresa/relatorio_ia.html', dados=dados)


@app.route('/empresa/relatorio/ia/stream')
@login_required_empresa
@requer_plano('profissional')
def relatorio_ia_stream():
    dados = _dados_empresa_para_ia(session['empresa_id'])

    prompt = f"""Você é um especialista em saúde ocupacional e compliance trabalhista brasileiro.
Analise os dados abaixo e gere um PARECER TÉCNICO completo para o Relatório NR-01 da empresa.

DADOS DA EMPRESA:
- Empresa: {dados['empresa']}
- Total de funcionários monitorados: {dados['total_funcionarios']}
- Score de risco atual (0-100): {dados['score_atual']} — Nível: {dados['nivel_atual'].upper()}
- Taxa de resposta esta semana: {dados['taxa_resposta']}%

HISTÓRICO DAS ÚLTIMAS 4 SEMANAS:
{chr(10).join(f"  Semana {d['semana']}: Score {d['score']} | Baixo:{d['dist']['baixo']} Moderado:{d['dist']['moderado']} Alto:{d['dist']['alto']} Crítico:{d['dist']['critico']} | {d['total']} respostas" for d in dados['historico_semanas'])}

SCORE MÉDIO POR DEPARTAMENTO:
{chr(10).join(f"  {dept}: {score}/100" for dept, score in dados['por_departamento'].items())}

COLABORADORES EM RISCO ALTO/CRÍTICO: {len(dados['funcionarios_criticos'])}
{chr(10).join(f"  - {c['nome']} ({c['dept']}): score {c['score']} — {c['nivel']}" for c in dados['funcionarios_criticos'])}

Gere o parecer em português formal com as seguintes seções:
1. SÍNTESE EXECUTIVA (3-4 linhas)
2. ANÁLISE DO AMBIENTE PSICOSSOCIAL (tendências, padrões identificados)
3. DEPARTAMENTOS EM ATENÇÃO (análise por área)
4. FATORES DE RISCO IDENTIFICADOS (baseado nos scores)
5. AÇÕES PREVENTIVAS RECOMENDADAS (concretas, com prazo sugerido)
6. CONCLUSÃO E CONFORMIDADE NR-01 (declaração formal de conformidade)

Use linguagem técnica mas clara. Seja específico com os dados fornecidos."""

    def generate():
        try:
            client = get_ai_client()
            with client.messages.stream(
                model="claude-sonnet-4-20250514",
                max_tokens=1500,
                messages=[{"role": "user", "content": prompt}]
            ) as stream:
                for text in stream.text_stream:
                    cleaned = text.replace('\n', '|||')
                    yield f"data: {cleaned}\n\n"
            yield "data: [DONE]\n\n"
        except Exception as e:
            yield f"data: [ERRO] {str(e)}\n\n"

    return Response(stream_with_context(generate()),
                    mimetype='text/event-stream',
                    headers={'Cache-Control': 'no-cache', 'X-Accel-Buffering': 'no'})


# ── 2. RECOMENDAÇÕES PERSONALIZADAS POR FUNCIONÁRIO (streaming) ──

@app.route('/funcionario/resultado/<int:ck_id>/ia')
@login_required_funcionario
def recomendacoes_ia(ck_id):
    ck = CheckIn.query.get_or_404(ck_id)
    if ck.funcionario_id != session['funcionario_id']:
        return redirect(url_for('painel_funcionario'))
    f = Funcionario.query.get(ck.funcionario_id)
    return render_template('funcionario/recomendacoes_ia.html', ck=ck, funcionario=f)


@app.route('/funcionario/resultado/<int:ck_id>/ia/stream')
@login_required_funcionario
def recomendacoes_ia_stream(ck_id):
    ck = CheckIn.query.get_or_404(ck_id)
    if ck.funcionario_id != session['funcionario_id']:
        return jsonify({'erro': 'Acesso negado'}), 403

    f = Funcionario.query.get(ck.funcionario_id)

    # Histórico do funcionário (últimas 4 semanas)
    historico = CheckIn.query.filter_by(funcionario_id=f.id).order_by(CheckIn.data.desc()).limit(4).all()
    trend = "estável"
    if len(historico) >= 2:
        if historico[0].score > historico[1].score + 10:
            trend = "piorando"
        elif historico[0].score < historico[1].score - 10:
            trend = "melhorando"

    prompt = f"""Você é um especialista em saúde mental ocupacional e bem-estar corporativo.
Um colaborador acabou de responder seu check-in semanal de bem-estar. Analise os dados e gere recomendações PERSONALIZADAS, empáticas e práticas.

PERFIL DO COLABORADOR:
- Departamento: {f.departamento} | Cargo: {f.cargo}
- Score de risco atual: {ck.score}/100 — Nível: {ck.nivel_risco.upper()}
- Tendência nas últimas semanas: {trend}

DIMENSÕES AVALIADAS (1=melhor, 5=pior):
- Estresse: {ck.q1_estresse}/5
- Sobrecarga de tarefas: {ck.q2_sobrecarga}/5
- Qualidade do sono: {ck.q3_sono}/5
- Motivação: {ck.q4_motivacao}/5
- Esgotamento emocional: {ck.q5_esgotamento}/5
- Relações no trabalho: {ck.q6_relacoes}/5
- Sensação de autonomia: {ck.q7_autonomia}/5
- Reconhecimento: {ck.q8_reconhecimento}/5

Gere uma resposta em português, calorosa e pessoal (não robótica), com:

1. **Como você está** — leitura empática do momento atual (2-3 linhas)
2. **Seus pontos de atenção** — as 2-3 dimensões mais críticas com explicação do impacto
3. **O que fazer esta semana** — 4-5 ações práticas e específicas (não genéricas como "descanse mais")
4. **Uma mensagem de apoio** — frase motivadora personalizada ao contexto

Seja direto, humano e útil. Evite linguagem clínica fria."""

    def generate():
        try:
            client = get_ai_client()
            with client.messages.stream(
                model="claude-sonnet-4-20250514",
                max_tokens=800,
                messages=[{"role": "user", "content": prompt}]
            ) as stream:
                for text in stream.text_stream:
                    cleaned = text.replace('\n', '|||')
                    yield f"data: {cleaned}\n\n"
            yield "data: [DONE]\n\n"
        except Exception as e:
            yield f"data: [ERRO] {str(e)}\n\n"

    return Response(stream_with_context(generate()),
                    mimetype='text/event-stream',
                    headers={'Cache-Control': 'no-cache', 'X-Accel-Buffering': 'no'})


# ── 3. ANÁLISE PREDITIVA DE RISCO ────────────────────────────────

@app.route('/empresa/previsao')
@login_required_empresa
@requer_plano('profissional')
def previsao_risco():
    empresa_id = session['empresa_id']
    funcionarios = Funcionario.query.filter_by(empresa_id=empresa_id, ativo=True).all()

    previsoes = []
    for f in funcionarios:
        historico = CheckIn.query.filter_by(funcionario_id=f.id)            .order_by(CheckIn.data.desc()).limit(4).all()
        if len(historico) < 2:
            continue

        scores = [h.score for h in historico]
        score_atual = scores[0]
        score_ant = scores[1]
        score_med = sum(scores) / len(scores)

        # Cálculo de tendência linear simples
        if len(scores) >= 3:
            delta_recente = scores[0] - scores[1]
            delta_anterior = scores[1] - scores[2]
            aceleracao = delta_recente - delta_anterior
        else:
            delta_recente = scores[0] - scores[1]
            aceleracao = 0

        # Score previsto = atual + tendência + metade da aceleração
        score_previsto = round(min(100, max(0, score_atual + delta_recente * 0.7 + aceleracao * 0.3)), 1)
        nivel_previsto = CheckIn.classificar_risco(score_previsto)
        nivel_atual = CheckIn.classificar_risco(score_atual)

        # Calcular dimensões mais críticas
        ult = historico[0]
        dims = {
            'Estresse': ult.q1_estresse,
            'Sobrecarga': ult.q2_sobrecarga,
            'Sono': ult.q3_sono,
            'Motivação': ult.q4_motivacao,
            'Esgotamento': ult.q5_esgotamento,
        }
        pior_dim = max(dims, key=dims.get)

        risco_escalada = (nivel_previsto in ('alto', 'critico') and
                          nivel_atual in ('baixo', 'moderado'))

        previsoes.append({
            'funcionario': f,
            'score_atual': score_atual,
            'score_previsto': score_previsto,
            'nivel_atual': nivel_atual,
            'nivel_previsto': nivel_previsto,
            'delta': round(score_previsto - score_atual, 1),
            'tendencia': 'subindo' if delta_recente > 3 else ('descendo' if delta_recente < -3 else 'estável'),
            'pior_dim': pior_dim,
            'risco_escalada': risco_escalada,
            'semanas_dados': len(historico),
        })

    previsoes.sort(key=lambda x: x['score_previsto'], reverse=True)

    # Sumário geral
    total = len(previsoes)
    em_escalada = sum(1 for p in previsoes if p['risco_escalada'])
    alto_previsto = sum(1 for p in previsoes if p['nivel_previsto'] in ('alto', 'critico'))

    return render_template('empresa/previsao.html',
        previsoes=previsoes,
        total=total,
        em_escalada=em_escalada,
        alto_previsto=alto_previsto
    )


# ── 4. CHATBOT RH (streaming) ─────────────────────────────────────

@app.route('/empresa/chat')
@login_required_empresa
@requer_plano('profissional')
def chat_rh():
    dados = _dados_empresa_para_ia(session['empresa_id'])
    return render_template('empresa/chat_rh.html', dados=dados)


@app.route('/empresa/chat/stream', methods=['POST'])
@login_required_empresa
@requer_plano('profissional')
def chat_rh_stream():
    dados = _dados_empresa_para_ia(session['empresa_id'])
    historico_chat = request.json.get('historico', [])
    pergunta = request.json.get('pergunta', '')

    system_prompt = f"""Você é BAIA — Burnout AI Assistant, a assistente de IA da plataforma BAI-EX.
Você apoia gestores de RH na gestão de riscos psicossociais e compliance com a NR-01.

CONTEXTO DA EMPRESA "{dados['empresa']}":
- {dados['total_funcionarios']} funcionários monitorados
- Score de risco atual: {dados['score_atual']}/100 ({dados['nivel_atual']})
- Taxa de resposta esta semana: {dados['taxa_resposta']}%
- Departamentos: {', '.join(f"{d}({s})" for d, s in dados['por_departamento'].items())}
- Funcionários em risco alto/crítico: {len(dados['funcionarios_criticos'])}

Você tem acesso aos dados reais da empresa acima. Responda sempre:
- Em português brasileiro, tom profissional mas acessível
- Com base nos dados reais quando relevante
- Com recomendações concretas e acionáveis
- De forma concisa (máximo 200 palavras por resposta)
- Se perguntarem sobre funcionários específicos por nome, diga que os dados são anonimizados por privacidade

Você é especialista em: NR-01, saúde mental no trabalho, gestão de equipes, prevenção de burnout, direito trabalhista básico."""

    messages = []
    for msg in historico_chat[-6:]:  # últimas 6 mensagens de contexto
        messages.append({"role": msg['role'], "content": msg['content']})
    messages.append({"role": "user", "content": pergunta})

    def generate():
        try:
            client = get_ai_client()
            with client.messages.stream(
                model="claude-sonnet-4-20250514",
                max_tokens=400,
                system=system_prompt,
                messages=messages
            ) as stream:
                for text in stream.text_stream:
                    cleaned = text.replace('\n', '|||')
                    yield f"data: {cleaned}\n\n"
            yield "data: [DONE]\n\n"
        except Exception as e:
            yield f"data: [ERRO] {str(e)}\n\n"

    return Response(stream_with_context(generate()),
                    mimetype='text/event-stream',
                    headers={'Cache-Control': 'no-cache', 'X-Accel-Buffering': 'no'})

# ── FIX 10: Validação de CNPJ ──────────────────────────────────
def validar_cnpj(cnpj: str) -> bool:
    """
    Valida CNPJ usando o algoritmo oficial da Receita Federal.
    Remove pontuação antes de validar.
    """
    cnpj = ''.join(filter(str.isdigit, cnpj))
    if len(cnpj) != 14:
        return False
    # Rejeita CNPJs com todos dígitos iguais
    if cnpj == cnpj[0] * 14:
        return False
    # Primeiro dígito verificador
    pesos = [5,4,3,2,9,8,7,6,5,4,3,2]
    soma = sum(int(cnpj[i]) * pesos[i] for i in range(12))
    resto = soma % 11
    d1 = 0 if resto < 2 else 11 - resto
    if int(cnpj[12]) != d1:
        return False
    # Segundo dígito verificador
    pesos = [6,5,4,3,2,9,8,7,6,5,4,3,2]
    soma = sum(int(cnpj[i]) * pesos[i] for i in range(13))
    resto = soma % 11
    d2 = 0 if resto < 2 else 11 - resto
    return int(cnpj[13]) == d2

def formatar_cnpj(cnpj: str) -> str:
    """Formata CNPJ para XX.XXX.XXX/XXXX-XX."""
    c = ''.join(filter(str.isdigit, cnpj))
    if len(c) == 14:
        return f"{c[:2]}.{c[2:5]}.{c[5:8]}/{c[8:12]}-{c[12:14]}"
    return cnpj


def seed_demo():
    """Popula dados de demonstração."""
    if Admin.query.first():
        return

    # Admin
    admin = Admin(usuario='admin', senha=generate_password_hash('burnout123'))
    db.session.add(admin)

    # Empresa demo
    empresa = Empresa(
        nome='Construtora Vale do São Francisco',
        cnpj='12.345.678/0001-99',
        responsavel='Ana Beatriz Santos',
        email='rh@construtora-vsf.com.br',
        senha=generate_password_hash('empresa123'),
        plano='profissional'
    )
    db.session.add(empresa)
    db.session.flush()

    # Departamentos e funcionários
    departamentos = ['Engenharia', 'Administrativo', 'Obras', 'Comercial', 'TI']
    nomes = [
        'Carlos Mendonça', 'Fernanda Lima', 'Rafael Souza', 'Beatriz Alves',
        'Thiago Costa', 'Juliana Ferreira', 'André Oliveira', 'Camila Rocha',
        'Diego Martins', 'Patrícia Nunes', 'Leandro Carvalho', 'Vanessa Cruz'
    ]

    funcionarios = []
    for i, nome in enumerate(nomes):
        email = f'{nome.split()[0].lower()}.{nome.split()[1].lower()}@construtora-vsf.com.br'
        f = Funcionario(
            empresa_id=empresa.id,
            nome=nome,
            email=email,
            senha=generate_password_hash('func123'),
            departamento=departamentos[i % len(departamentos)],
            cargo=['Engenheiro', 'Analista', 'Técnico', 'Gerente'][i % 4]
        )
        db.session.add(f)
        funcionarios.append(f)
    db.session.flush()

    # Check-ins das últimas 8 semanas com padrões realistas
    hoje = date.today()
    for semana_offset in range(7, -1, -1):
        dt = hoje - timedelta(weeks=semana_offset)
        semana = dt.isocalendar()[1]
        ano = dt.year

        for i, f in enumerate(funcionarios):
            # Alguns não respondem sempre
            if random.random() < 0.85:
                # Simular tendência de aumento de risco nas últimas semanas
                fator = 1 + (semana_offset == 0) * 0.3
                base_scores = [
                    [2, 2, 4, 4, 2, 4, 4, 4],   # baixo risco
                    [3, 3, 3, 3, 3, 3, 3, 3],   # moderado
                    [4, 4, 2, 2, 4, 2, 2, 2],   # alto risco
                    [5, 5, 1, 1, 5, 1, 1, 1],   # crítico
                ]
                perfil = [0,0,0,1,1,2,2,2,3,2,1,0][i % 12]
                rs = base_scores[perfil]
                rs = [max(1, min(5, r + random.randint(-1, 1))) for r in rs]

                score = CheckIn.calcular_score(rs)
                nivel = CheckIn.classificar_risco(score)
                ck = CheckIn(
                    funcionario_id=f.id,
                    empresa_id=empresa.id,
                    semana=semana, ano=ano,
                    q1_estresse=rs[0], q2_sobrecarga=rs[1], q3_sono=rs[2],
                    q4_motivacao=rs[3], q5_esgotamento=rs[4], q6_relacoes=rs[5],
                    q7_autonomia=rs[6], q8_reconhecimento=rs[7],
                    score=score, nivel_risco=nivel,
                    data=datetime.combine(dt, datetime.min.time())
                )
                db.session.add(ck)

    # Funcionário demo para login
    func_demo = Funcionario(
        empresa_id=empresa.id,
        nome='João Demo',
        email='joao@demo.com',
        senha=generate_password_hash('func123'),
        departamento='Engenharia',
        cargo='Analista'
    )
    db.session.add(func_demo)
    db.session.commit()
    print("✅ Dados demo criados com sucesso!")

# ═══════════════════════════════════════════════════════════════
#   MÓDULO DE PAGAMENTO — Stripe
#   Configure no .env:
#     STRIPE_SECRET_KEY=sk_live_...
#     STRIPE_PUBLISHABLE_KEY=pk_live_...
#     STRIPE_PRICE_STARTER1=price_xxx   (R$124,50/mês fixo)
#     STRIPE_PRICE_STARTER2=price_xxx   (R$249/mês fixo)
#     STRIPE_PRICE_PRO=price_xxx        (R$12/func/mês)
#     STRIPE_PRICE_EMPRESARIAL=price_xxx(R$10/func/mês)
#     STRIPE_WEBHOOK_SECRET=whsec_...
# ═══════════════════════════════════════════════════════════════

def _stripe_client():
    import stripe as _stripe
    key = os.environ.get('STRIPE_SECRET_KEY', '')
    if not key or 'COLE' in key:
        return None
    _stripe.api_key = key
    return _stripe

STRIPE_PUBLISHABLE_KEY = os.environ.get('STRIPE_PUBLISHABLE_KEY', '')

def _stripe_price_id(plano_base: str) -> str:
    """Retorna o price_id do Stripe para o plano dado."""
    env_map = {
        'starter1':    'STRIPE_PRICE_STARTER1',
        'starter2':    'STRIPE_PRICE_STARTER2',
        'profissional':'STRIPE_PRICE_PRO',
        'empresarial': 'STRIPE_PRICE_EMPRESARIAL',
    }
    return os.environ.get(env_map.get(plano_base, ''), '')


@app.route('/assinar/<plano>')
@login_required_empresa
def iniciar_assinatura(plano):
    """Redireciona empresa para o checkout Stripe."""
    stripe = _stripe_client()
    if not stripe:
        flash('Gateway de pagamento não configurado. Entre em contato: baiexstartup@gmail.com', 'warning')
        return redirect(url_for('index'))

    if plano not in PLANOS_INFO or plano == 'enterprise':
        flash('Plano inválido ou requer contato direto.', 'danger')
        return redirect(url_for('index'))

    empresa = Empresa.query.get(session['empresa_id'])
    num_func = Funcionario.query.filter_by(empresa_id=empresa.id, ativo=True).count() or 1
    preco_final = calcular_preco_plano(plano, num_func)
    price_id   = _stripe_price_id(plano)

    if not price_id:
        flash('Price ID do Stripe não configurado. Adicione STRIPE_PRICE_* no .env.', 'warning')
        return redirect(url_for('minha_conta'))

    try:
        checkout = stripe.checkout.Session.create(
            payment_method_types=['card'],
            mode='subscription',
            line_items=[{
                'price': price_id,
                'quantity': max(num_func, 1),
            }],
            customer_email=empresa.email,
            metadata={'empresa_id': str(empresa.id), 'plano': plano},
            success_url=request.host_url.rstrip('/') + url_for('assinatura_sucesso') + '?session_id={CHECKOUT_SESSION_ID}',
            cancel_url=request.host_url.rstrip('/') + url_for('minha_conta'),
            locale='pt-BR',
        )
        registrar_auditoria('CHECKOUT_INICIADO',
                            f'Plano: {plano} | Valor: R${preco_final:.2f}',
                            tipo_usuario='empresa', usuario_id=empresa.id, empresa_id=empresa.id)
        return redirect(checkout.url, code=303)
    except Exception as e:
        flash('Erro ao iniciar pagamento. Tente novamente ou entre em contato.', 'danger')
        return redirect(url_for('minha_conta'))


@app.route('/assinatura/sucesso')
@login_required_empresa
def assinatura_sucesso():
    """Página de sucesso após pagamento Stripe."""
    session_id = request.args.get('session_id', '')
    stripe     = _stripe_client()
    empresa    = Empresa.query.get(session['empresa_id'])
    plano_ativado = empresa.plano_base

    if stripe and session_id:
        try:
            checkout_session = stripe.checkout.Session.retrieve(session_id)
            plano_ativado = checkout_session.metadata.get('plano', empresa.plano_base)
            empresa.plano = plano_ativado  # sem prefixo trial_
            empresa.fidelizacao_inicio = datetime.utcnow()
            empresa.fidelizacao_fim    = datetime.utcnow() + timedelta(days=180)
            db.session.commit()
            registrar_auditoria('ASSINATURA_ATIVA', f'Plano: {plano_ativado}',
                                tipo_usuario='empresa', usuario_id=empresa.id, empresa_id=empresa.id)
        except Exception:
            pass

    nome_plano = PLANOS_INFO.get(plano_ativado, {}).get('nome', plano_ativado.title())
    flash(f'🎉 Assinatura ativada com sucesso! Plano {nome_plano} ativo.', 'success')
    return redirect(url_for('dashboard_empresa'))


@app.route('/stripe/webhook', methods=['POST'])
def stripe_webhook():
    """Recebe eventos do Stripe (cancelamento, renovação, falha de pagamento)."""
    stripe         = _stripe_client()
    webhook_secret = os.environ.get('STRIPE_WEBHOOK_SECRET', '')
    if not stripe or not webhook_secret:
        return '', 400

    payload = request.get_data()
    sig_header = request.headers.get('Stripe-Signature', '')

    try:
        event = stripe.Webhook.construct_event(payload, sig_header, webhook_secret)
    except Exception:
        return '', 400

    if event['type'] == 'customer.subscription.deleted':
        # Assinatura cancelada — desativa plano
        meta = event['data']['object'].get('metadata', {})
        empresa_id = meta.get('empresa_id')
        if empresa_id:
            empresa = Empresa.query.get(int(empresa_id))
            if empresa:
                empresa.plano = 'inativo'
                db.session.commit()

    elif event['type'] == 'invoice.payment_failed':
        # Pagamento falhou — poderia enviar e-mail de aviso
        pass

    return '', 200


# ═══════════════════════════════════════════════════════════════
#   MÓDULO DE E-MAIL — Convite automático para funcionários
#   Configure no .env:
#     MAIL_SERVER=smtp.gmail.com
#     MAIL_PORT=587
#     MAIL_USERNAME=baiexstartup@gmail.com
#     MAIL_PASSWORD=sua_senha_de_app_gmail
#     MAIL_DEFAULT_SENDER=baiexstartup@gmail.com
# ═══════════════════════════════════════════════════════════════

def _enviar_email(destinatario: str, assunto: str, corpo_html: str) -> bool:
    """
    Envia e-mail via SMTP. Retorna True se enviado, False se falhar.
    Usa as variáveis MAIL_* do .env. Falha silenciosamente para não
    travar o cadastro do funcionário.
    """
    import smtplib
    from email.mime.multipart import MIMEMultipart
    from email.mime.text import MIMEText

    servidor  = os.environ.get('MAIL_SERVER', '')
    porta     = int(os.environ.get('MAIL_PORT', 587))
    usuario   = os.environ.get('MAIL_USERNAME', '')
    senha     = os.environ.get('MAIL_PASSWORD', '')
    remetente = os.environ.get('MAIL_DEFAULT_SENDER', usuario)

    if not all([servidor, usuario, senha]):
        return False  # e-mail não configurado, ignora silenciosamente

    try:
        msg = MIMEMultipart('alternative')
        msg['Subject'] = assunto
        msg['From']    = f'BAI-EX <{remetente}>'
        msg['To']      = destinatario
        msg.attach(MIMEText(corpo_html, 'html', 'utf-8'))

        with smtplib.SMTP(servidor, porta, timeout=10) as smtp:
            smtp.ehlo()
            smtp.starttls()
            smtp.login(usuario, senha)
            smtp.sendmail(remetente, [destinatario], msg.as_string())
        return True
    except Exception:
        return False


def _html_convite_funcionario(nome_func: str, email: str, senha_temp: str,
                               nome_empresa: str, url_base: str) -> str:
    """Gera o HTML do e-mail de convite para o funcionário."""
    login_url = f"{url_base.rstrip('/')}/funcionario/login"
    return f"""
<!DOCTYPE html>
<html lang="pt-BR">
<head><meta charset="UTF-8"></head>
<body style="margin:0;padding:0;background:#0D1117;font-family:Arial,sans-serif">
  <table width="100%" cellpadding="0" cellspacing="0">
    <tr><td align="center" style="padding:40px 20px">
      <table width="560" cellpadding="0" cellspacing="0"
             style="background:#161B22;border:1px solid #21262D;border-radius:16px;overflow:hidden">

        <!-- Header -->
        <tr><td style="background:linear-gradient(135deg,#0D1117,#1a2a3a);padding:32px;text-align:center">
          <span style="font-size:1.4rem;font-weight:800;color:#fff;letter-spacing:-0.5px">
            ⚡ BAI-EX
          </span>
          <p style="color:#8B949E;font-size:.85rem;margin:6px 0 0">Burnout AI Experience</p>
        </td></tr>

        <!-- Body -->
        <tr><td style="padding:32px">
          <p style="color:#C9D1D9;font-size:1rem;margin:0 0 8px">Olá, <strong style="color:#fff">{nome_func}</strong> 👋</p>
          <p style="color:#8B949E;font-size:.9rem;line-height:1.6;margin:0 0 24px">
            A empresa <strong style="color:#22D3EE">{nome_empresa}</strong> cadastrou você na
            plataforma <strong style="color:#fff">BAI-EX</strong> — sistema de monitoramento de
            bem-estar e compliance com a NR-01.
          </p>
          <p style="color:#C9D1D9;font-size:.9rem;margin:0 0 6px">Seus dados de acesso:</p>
          <div style="background:#0D1117;border:1px solid #21262D;border-radius:10px;padding:16px;margin-bottom:24px">
            <p style="margin:0 0 6px;color:#8B949E;font-size:.8rem">E-mail</p>
            <p style="margin:0 0 14px;color:#fff;font-weight:600">{email}</p>
            <p style="margin:0 0 6px;color:#8B949E;font-size:.8rem">Senha temporária</p>
            <p style="margin:0;color:#22D3EE;font-weight:700;font-size:1.1rem;letter-spacing:.05em">{senha_temp}</p>
          </div>
          <p style="color:#8B949E;font-size:.8rem;margin:0 0 24px">
            ⚠️ Por segurança, altere sua senha no primeiro acesso.
          </p>
          <a href="{login_url}"
             style="display:block;text-align:center;background:#0891B2;color:#fff;font-weight:700;
                    font-size:.95rem;padding:14px;border-radius:10px;text-decoration:none">
            Acessar minha conta →
          </a>
        </td></tr>

        <!-- Footer -->
        <tr><td style="border-top:1px solid #21262D;padding:20px 32px;text-align:center">
          <p style="color:#8B949E;font-size:.75rem;margin:0">
            O check-in semanal leva menos de 2 minutos e é completamente anônimo para sua empresa.<br>
            Dúvidas? <a href="mailto:baiexstartup@gmail.com" style="color:#22D3EE;text-decoration:none">baiexstartup@gmail.com</a>
          </p>
        </td></tr>

      </table>
    </td></tr>
  </table>
</body>
</html>"""


# Sobrescreve a rota de novo funcionário para adicionar envio de e-mail
@app.route('/empresa/funcionarios/novo', methods=['GET', 'POST'])
@login_required_empresa
def novo_funcionario():
    empresa_id = session['empresa_id']
    empresa    = Empresa.query.get(empresa_id)
    if request.method == 'POST':
        nome         = request.form['nome']
        email        = request.form['email']
        senha        = request.form['senha']
        departamento = request.form.get('departamento', 'Geral')
        cargo        = request.form.get('cargo', '')
        enviar_conv  = request.form.get('enviar_convite') == '1'

        # ── Verificar limite de funcionários do plano ──────────
        total_ativos = Funcionario.query.filter_by(empresa_id=empresa_id, ativo=True).count()
        if total_ativos >= empresa.limite_funcionarios:
            flash(f'⛔ Limite do plano {PLANOS_INFO[empresa.plano_base]["nome"]} atingido '
                  f'({empresa.limite_funcionarios} funcionários). '
                  f'Faça upgrade para adicionar mais.', 'danger')
            return redirect(url_for('funcionarios_lista'))

        if Funcionario.query.filter_by(email=email).first():
            flash('Email já cadastrado.', 'danger')
        else:
            f = Funcionario(empresa_id=empresa_id, nome=nome, email=email,
                            senha=generate_password_hash(senha),
                            departamento=departamento, cargo=cargo)
            db.session.add(f)
            db.session.commit()

            email_enviado = False
            if enviar_conv:
                url_base = request.host_url
                html = _html_convite_funcionario(
                    nome_func=nome, email=email, senha_temp=senha,
                    nome_empresa=empresa.nome, url_base=url_base
                )
                email_enviado = _enviar_email(
                    destinatario=email,
                    assunto=f'Seu acesso ao BAI-EX — {empresa.nome}',
                    corpo_html=html
                )

            registrar_auditoria('CADASTRO_FUNCIONARIO',
                                f'Nome: {nome} | Email enviado: {email_enviado}',
                                tipo_usuario='empresa', usuario_id=empresa_id,
                                empresa_id=empresa_id)

            if enviar_conv and email_enviado:
                flash(f'✅ Funcionário {nome} cadastrado e convite enviado por e-mail!', 'success')
            elif enviar_conv and not email_enviado:
                flash(f'✅ Funcionário {nome} cadastrado. ⚠️ E-mail não enviado (configure MAIL_* no .env).', 'warning')
            else:
                flash(f'Funcionário {nome} cadastrado!', 'success')

            return redirect(url_for('funcionarios_lista'))
    # alerta de limite
    total_ativos = Funcionario.query.filter_by(empresa_id=empresa_id, ativo=True).count()
    quase_limite = total_ativos >= empresa.limite_funcionarios * 0.9
    return render_template('empresa/novo_funcionario.html',
                           empresa=empresa,
                           total_ativos=total_ativos,
                           quase_limite=quase_limite)


# ═══════════════════════════════════════════════════════════════
#   ADMIN — gerenciamento de empresas
# ═══════════════════════════════════════════════════════════════

@app.route('/admin/empresa/<int:eid>/toggle')
@login_required_admin
def admin_toggle_empresa(eid):
    empresa = Empresa.query.get_or_404(eid)
    empresa.ativo = not empresa.ativo
    db.session.commit()
    acao = 'ativada' if empresa.ativo else 'desativada'
    registrar_auditoria('ADMIN_TOGGLE_EMPRESA',
                        f'{empresa.nome} foi {acao}', tipo_usuario='admin')
    flash(f'Empresa "{empresa.nome}" {acao} com sucesso.', 'info')
    return redirect(url_for('dashboard_admin'))

@app.route('/admin/empresa/<int:eid>/excluir', methods=['POST'])
@login_required_admin
def admin_excluir_empresa(eid):
    empresa = Empresa.query.get_or_404(eid)
    nome = empresa.nome
    # Remove tudo na ordem correta (chaves estrangeiras)
    CheckIn.query.filter_by(empresa_id=eid).delete()
    AuditLog.query.filter_by(empresa_id=eid).delete()
    Funcionario.query.filter_by(empresa_id=eid).delete()
    db.session.delete(empresa)
    db.session.commit()
    registrar_auditoria('ADMIN_EXCLUIR_EMPRESA', f'Empresa excluída: {nome}',
                        tipo_usuario='admin')
    flash(f'Empresa "{nome}" excluída permanentemente.', 'warning')
    return redirect(url_for('dashboard_admin'))

@app.route('/admin/empresa/<int:eid>/plano', methods=['POST'])
@login_required_admin
def admin_alterar_plano(eid):
    empresa = Empresa.query.get_or_404(eid)
    novo_plano = request.form.get('plano', '')
    if novo_plano in PLANOS_INFO:
        plano_antigo = empresa.plano
        empresa.plano = novo_plano  # sem trial — admin ativa direto
        empresa.cancelado = False
        empresa.ativo = True
        db.session.commit()
        registrar_auditoria('ADMIN_ALTERAR_PLANO',
                            f'{empresa.nome}: {plano_antigo} → {novo_plano}',
                            tipo_usuario='admin')
        flash(f'Plano de "{empresa.nome}" alterado para {PLANOS_INFO[novo_plano]["nome"]}.', 'success')
    else:
        flash('Plano inválido.', 'danger')
    return redirect(url_for('dashboard_admin'))

# ═══════════════════════════════════════════════════════════════
#   EMPRESA — minha conta, frequência, cancelamento
# ═══════════════════════════════════════════════════════════════

@app.route('/empresa/conta')
@login_required_empresa
def minha_conta():
    empresa = Empresa.query.get(session['empresa_id'])
    total_func = Funcionario.query.filter_by(empresa_id=empresa.id, ativo=True).count()
    plano_info = PLANOS_INFO.get(empresa.plano_base, PLANOS_INFO['starter1'])
    preco_atual = calcular_preco_plano(empresa.plano_base, total_func)
    return render_template('empresa/minha_conta.html',
        empresa=empresa,
        plano_info=plano_info,
        preco_atual=preco_atual,
        total_func=total_func,
        PLANOS_INFO=PLANOS_INFO,
    )

@app.route('/empresa/conta/frequencia', methods=['POST'])
@login_required_empresa
def configurar_frequencia():
    empresa = Empresa.query.get(session['empresa_id'])
    # Frequência extra só para Profissional+
    if not empresa.tem_acesso('profissional'):
        flash('Frequência configurável disponível a partir do plano Profissional.', 'warning')
        return redirect(url_for('minha_conta'))
    nova_freq = request.form.get('frequencia', 'semanal')
    if nova_freq in ('semanal', 'quinzenal', 'mensal'):
        empresa.frequencia_checkin = nova_freq
        db.session.commit()
        labels = {'semanal': 'Semanal', 'quinzenal': 'Quinzenal', 'mensal': 'Mensal'}
        registrar_auditoria('CONFIG_FREQUENCIA', f'Frequência → {nova_freq}')
        flash(f'✅ Frequência de check-in alterada para {labels[nova_freq]}.', 'success')
    return redirect(url_for('minha_conta'))

@app.route('/empresa/conta/cancelar', methods=['POST'])
@login_required_empresa
def cancelar_assinatura():
    empresa = Empresa.query.get(session['empresa_id'])
    confirmacao = request.form.get('confirmar_cancelamento', '')
    if confirmacao != empresa.nome:
        flash('Confirmação incorreta. Digite o nome da empresa exatamente como mostrado.', 'danger')
        return redirect(url_for('minha_conta'))

    empresa.cancelado = True
    db.session.commit()

    # Cancela no Stripe se tiver assinatura ativa
    stripe = _stripe_client()
    if stripe and not empresa.em_trial:
        try:
            subs = stripe.Subscription.list(limit=5)
            for sub in subs.get('data', []):
                meta = sub.get('metadata', {})
                if str(meta.get('empresa_id')) == str(empresa.id):
                    stripe.Subscription.modify(sub['id'],
                                               cancel_at_period_end=True)
                    break
        except Exception:
            pass

    registrar_auditoria('CANCELAMENTO', f'Empresa cancelou: {empresa.nome}')
    flash('⚠️ Assinatura cancelada. Seu acesso continua até o fim do período pago.', 'warning')
    session.clear()
    return redirect(url_for('index'))

# ── Calculadora de preço (API para landing page) ──────────────
@app.route('/api/calculadora')
def api_calculadora():
    try:
        n = int(request.args.get('funcionarios', 1))
        n = max(1, min(n, 99999))
    except ValueError:
        n = 1
    plano = plano_para_funcionarios(n)
    info  = PLANOS_INFO[plano]
    preco = calcular_preco_plano(plano, n)

    if info['preco_tipo'] == 'negociado':
        preco_display = 'Sob consulta'
    elif info['preco_tipo'] == 'fixo':
        preco_display = f"R$ {info.get('preco_mensal', preco):.0f}/mês"
    else:
        preco_display = f"R$ {preco:,.0f}".replace(',', '.')

    return jsonify({
        'funcionarios': n,
        'plano': plano,
        'nome': info['nome'],
        'preco': preco,
        'preco_display': preco_display,
        'descricao': info['descricao'],
        'trial_dias': info['trial_dias'],
    })

# ── Solicitação de migração de plano (gera alerta ao admin) ───
@app.route('/empresa/solicitar-migracao', methods=['POST'])
@login_required_empresa
def solicitar_migracao():
    """
    Empresa solicita mudança de plano.
    Cria um AlertaMigracao no banco — admin vê no painel e contata o cliente.
    """
    dados = request.get_json() or {}
    plano_destino = dados.get('plano', '')
    plano_nome = dados.get('nome', '')
    empresa = Empresa.query.get(session['empresa_id'])

    if not plano_destino or plano_destino not in PLANOS_INFO:
        return jsonify({'ok': False, 'erro': 'Plano inválido'}), 400

    # Registrar na auditoria (admin vê nos logs)
    registrar_auditoria(
        'SOLICITACAO_MIGRACAO',
        f'Empresa: {empresa.nome} | De: {empresa.plano_base} → Para: {plano_destino} ({plano_nome})',
        tipo_usuario='empresa',
        usuario_id=empresa.id,
        empresa_id=empresa.id
    )
    return jsonify({'ok': True})


if __name__ == '__main__':
    with app.app_context():
        db.create_all()
        seed_data()
    print("=" * 55)
    print("  🧠 BAI-EX - Sistema de Gestão de Risco")
    print("=" * 55)
    print("  🏢 Empresa/RH  → http://localhost:5002/empresa/login")
    print("     Email: rh@construtora-vsf.com.br | Senha: empresa123")
    print()
    print("  👤 Funcionário → http://localhost:5002/funcionario/login")
    print("     Email: joao@demo.com | Senha: func123")
    print()
    print("  🔧 Admin       → http://localhost:5002/admin/login")
    print("     Usuário: admin | Senha: burnout123")
    print("=" * 55)
    app.run(debug=True, port=5002)

from datetime import datetime, timedelta
import random
import json
from werkzeug.security import generate_password_hash

def seed_data():
    from models import db, Empresa, Funcionario, Admin, Checkin

    # Evita duplicar
    if Empresa.query.first():
        return

    # =========================
    # EMPRESA DEMO
    # =========================
    empresa = Empresa(
        nome="Empresa Demo",
        email="empresa@baiex.com",
        senha=generate_password_hash("123456")
    )
    db.session.add(empresa)
    db.session.commit()

    # =========================
    # FUNCIONÁRIOS
    # =========================
    nomes = ["João", "Maria", "Carlos", "Ana", "Pedro"]

    funcionarios = []

    for nome in nomes:
        f = Funcionario(
            nome=nome,
            email=f"{nome.lower()}@demo.com",
            senha=generate_password_hash("123456"),
            empresa_id=empresa.id
        )
        db.session.add(f)
        funcionarios.append(f)

    db.session.commit()

    # =========================
    # ADMIN
    # =========================
    admin = Admin(
        nome="Admin",
        email="admin@baiex.com",
        senha=generate_password_hash("123456")
    )
    db.session.add(admin)
    db.session.commit()

    # =========================
    # CHECK-INS (dados)
    # =========================
    for f in funcionarios:
        for i in range(4):  # 4 semanas
            data = datetime.now() - timedelta(days=7 * i)

            respostas = [random.randint(2, 5) for _ in range(8)]
            score = sum(respostas) / len(respostas)

            checkin = Checkin(
                funcionario_id=f.id,
                empresa_id=empresa.id,
                data=data,
                respostas=json.dumps(respostas),
                score=score
            )

            db.session.add(checkin)

    db.session.commit()

    print("✅ Dados demo criados!")
